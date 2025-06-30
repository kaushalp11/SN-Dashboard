import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime, timezone, timedelta
import io
import plotly.io as pio
import numpy as np
from groq import Groq
import json

# --- Imports for DOCX Export ---
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# --- Global Constants ---
DAY_FIRST_TRUE_COLS = ['Opened', 'Resolved', 'Closed']
LLM_PROVIDER = "groq"
# --- Page Configuration ---
st.set_page_config(
    page_title="ServiceNow CSM Operational KPI Dashboard",
    page_icon="📊",
    layout="wide"
)

# --- LLM Helper Functions ---
@st.cache_data(ttl=3600, show_spinner=False)
def get_genai_summary(chart_title, chart_purpose, chart_benefit, chart_data_summary_str=None, chart_type="general"):
    if not chart_title and not chart_purpose:
        return "Cannot generate summary: Chart title and purpose are missing."
    base_prompt = f"""
    You are an expert data analyst providing insights for a ServiceNow performance dashboard.
    Analyze the following chart:

    Chart Title: {chart_title}
    Chart Purpose: {chart_purpose}
    Chart Benefit: {chart_benefit}
    """
    if chart_data_summary_str:
        base_prompt += f"\nKey Data Points/Summary: {chart_data_summary_str}\n"
    else:
        base_prompt += "\n(No specific data points provided, rely on title, purpose, and benefit for general analysis)\n"
    base_prompt += """
    Provide your analysis in a structured markdown format with the following sections:
    1.  **Chart Overview:** Briefly describe what this chart likely visualizes based on its title and purpose.
    2.  **Key Insights & Trends:** What are the most important patterns, anomalies, or significant observations a user should take away? If it's a trend chart, describe the trend.
    3.  **Potential Implications & Predictive Thoughts:** Based on the insights, what could be the implications? If trends are present, what might they suggest for the near future? (State any assumptions or if predictions are speculative due to limited data).
    4.  **Actionable Recommendations:** Suggest 1-2 concrete actions or areas for further investigation based on your analysis.

    Keep your analysis concise, insightful, and professional.
    """
    try:
        if LLM_PROVIDER == "groq":
            if "GROQ_API_KEY" not in st.secrets:
                return "Error: GROQ_API_KEY not found in secrets."
            client = Groq(api_key=st.secrets["GROQ_API_KEY"])
            chat_completion = client.chat.completions.create(
                messages=[{"role": "user", "content": base_prompt}],
                model="llama3-8b-8192", temperature=0.7, max_tokens=800,
            )
            return chat_completion.choices[0].message.content
        elif LLM_PROVIDER == "azure":
            st.warning("Azure OpenAI provider selected, but implementation is a placeholder.")
            return "Azure OpenAI summary (placeholder)."
        else:
            return "Error: Invalid LLM_PROVIDER configured."
    except Exception as e:
        st.error(f"Error calling LLM API: {e}")
        return f"Could not generate summary due to an API error: {str(e)}"

def generate_data_summary_for_llm(df_chart, chart_type, x_col=None, y_col=None, value_col=None, name_col=None, top_n=5):
    if df_chart is None or df_chart.empty:
        return "No data available for summary."
    summary_parts = []
    try:
        if chart_type == "bar" and x_col and y_col:
            summary_parts.append(f"Bar chart showing '{y_col}' per '{x_col}'.")
            if x_col in df_chart.columns and y_col in df_chart.columns and pd.api.types.is_numeric_dtype(df_chart[y_col]):
                top_items = df_chart.nlargest(top_n, y_col)
                for _, row in top_items.iterrows():
                    summary_parts.append(f"- {row[x_col]}: {row[y_col]:.2f}")
            else: # Fallback if columns or type are not as expected
                summary_parts.append(f"Top {top_n} rows: {df_chart.head(top_n).to_dict(orient='records')}")
        elif chart_type == "bar_grouped" and x_col and isinstance(y_col, list): # Specific for opened vs closed
            summary_parts.append(f"Grouped bar chart showing trends for {', '.join(y_col)} over '{x_col}'.")
            if x_col in df_chart.columns:
                for col in y_col:
                    if col in df_chart.columns and pd.api.types.is_numeric_dtype(df_chart[col]):
                        summary_parts.append(f"Average {col}: {df_chart[col].mean():.2f}")
                summary_parts.append(f"Total entries: {len(df_chart)}.")
                summary_parts.append(f"First period: {df_chart[x_col].min()}, Last period: {df_chart[x_col].max()}.")
            else:
                summary_parts.append(f"Data snapshot: {df_chart.head(top_n).to_dict(orient='records')}")
        elif chart_type == "pie" and name_col and value_col:
            summary_parts.append(f"Pie chart showing distribution of '{value_col}' by '{name_col}'.")
            if name_col in df_chart.columns and value_col in df_chart.columns and pd.api.types.is_numeric_dtype(df_chart[value_col]):
                top_items = df_chart.nlargest(top_n, value_col)
                total_value = df_chart[value_col].sum()
                for _, row in top_items.iterrows():
                    percentage = (row[value_col] / total_value * 100) if total_value else 0
                    summary_parts.append(f"- {row[name_col]}: {row[value_col]} ({percentage:.1f}%)")
            else:
                summary_parts.append(f"Top {top_n} rows: {df_chart.head(top_n).to_dict(orient='records')}")
        elif chart_type == "line" and x_col and y_col:
            summary_parts.append(f"Line chart showing trend of '{y_col}' over '{x_col}'.")
            if y_col in df_chart.columns and pd.api.types.is_numeric_dtype(df_chart[y_col]):
                summary_parts.append(f"Min {y_col}: {df_chart[y_col].min():.2f}, Max {y_col}: {df_chart[y_col].max():.2f}, Avg {y_col}: {df_chart[y_col].mean():.2f}.")
            else:
                summary_parts.append(f"Data snapshot: {df_chart.head(top_n).to_dict(orient='records')}")
        elif chart_type == "box" and x_col and y_col:
            summary_parts.append(f"Box plot showing distribution of '{y_col}' by '{x_col}'.")
            if x_col in df_chart.columns and y_col in df_chart.columns and pd.api.types.is_numeric_dtype(df_chart[y_col]):
                for group in df_chart[x_col].unique():
                    group_data = df_chart[df_chart[x_col] == group][y_col]
                    if not group_data.empty:
                        summary_parts.append(f"- {group}: Median {group_data.median():.2f}, Mean {group_data.mean():.2f}, Min {group_data.min():.2f}, Max {group_data.max():.2f}.")
            else:
                summary_parts.append(f"Data snapshot: {df_chart.head(top_n).to_dict(orient='records')}")
        elif chart_type == "heatmap" and x_col and y_col: # Note: value_col for heatmap is often implicit in pivot
            summary_parts.append(f"Heatmap showing ticket counts across '{y_col}' (rows) and '{x_col}' (columns).")
            if df_chart is not None and not df_chart.empty:
                max_val = df_chart.values.max() if df_chart.size > 0 and pd.api.types.is_numeric_dtype(df_chart.values) else 'N/A'
                min_val = df_chart.values.min() if df_chart.size > 0 and pd.api.types.is_numeric_dtype(df_chart.values) else 'N/A'
                summary_parts.append(f"Max ticket count: {max_val}, Min ticket count: {min_val}.")
                summary_parts.append(f"Dimensions: {len(df_chart.index)} {y_col} and {len(df_chart.columns)} {x_col} buckets.")
            else:
                summary_parts.append("No specific pivoted data for detailed heatmap summary.")
        else:
            summary_parts.append(f"General data with {len(df_chart)} rows and columns: {', '.join(df_chart.columns)}.")
        return " ".join(summary_parts)
    except Exception as e:
        return f"Could not summarize data: {e}"

# --- Helper Functions (load_data, get_time_periods, calculate_sla_percentage, calculate_kpis) ---
@st.cache_data
def load_data(uploaded_file):
    try:
        df = pd.read_excel(uploaded_file, engine='openpyxl')
        required_cols = ['Task', 'Opened', 'Has breached', 'Target', 'State', 'Priority', 'Region', 'Account', 'Resolved', 'Closed']
        optional_cols = ['Start time', 'Stop time', 'Original breach time', 'Assignment group', 'Category', 'Subcategory']
        date_columns_core = ['Opened', 'Resolved', 'Closed']
        date_columns_optional = ['Start time', 'Stop time', 'Original breach time']
        date_columns = date_columns_core + date_columns_optional

        missing_req = [col for col in required_cols if col not in df.columns]
        if missing_req:
            st.error(f"Error: Required column(s) missing: {', '.join(missing_req)}. Please check the file structure.")
            st.stop()
        missing_dates = [col for col in date_columns_core if col not in df.columns]
        if missing_dates:
            st.warning(f"Warning: Core date column(s) missing: {', '.join(missing_dates)}. Some calculations might be affected.")

        for col in date_columns:
            if col in df.columns:
                df[col] = pd.to_datetime(df[col], errors='coerce', dayfirst=(col in DAY_FIRST_TRUE_COLS))
        if 'Has breached' in df.columns:
            df['Has breached'] = df['Has breached'].map({
                'TRUE': True, True: True, 'True': True, 'true': True,
                'FALSE': False, False: False, 'False': False, 'false': False
            }).fillna(pd.NA)
            try: df['Has breached'] = df['Has breached'].astype('boolean')
            except TypeError: st.warning("Could not fully convert 'Has breached' column to boolean.")
        if 'Opened' in df.columns and pd.api.types.is_datetime64_any_dtype(df['Opened']):
             df['Opened_YearMonth'] = df['Opened'].dt.to_period('M')
        if 'Closed' in df.columns and pd.api.types.is_datetime64_any_dtype(df['Closed']):
            df['Closed_YearMonth'] = df['Closed'].dt.to_period('M')
        if 'Priority' in df.columns: df['Priority'] = df['Priority'].astype(str)
        for col in optional_cols:
            if col not in df.columns: st.info(f"Optional column '{col}' not found. Dependent charts may not be available.")
            elif col in ['Category', 'Subcategory', 'Assignment group'] and col in df.columns:
                 df[col] = df[col].astype(str).fillna('N/A')
        st.success("Data loaded and preprocessed successfully!")
        return df
    except Exception as e:
        st.error(f"Error loading or processing file: {e}"); st.exception(e); st.stop()

def get_time_periods(df, date_column, freq):
    if df is None or df.empty or date_column not in df.columns or df[date_column].isnull().all(): return ['N/A']
    try:
        period_freq = freq if freq != 'Y' else 'A-DEC'
        valid_dates = pd.to_datetime(df[date_column], errors='coerce', dayfirst=(date_column in DAY_FIRST_TRUE_COLS)).dropna()
        if valid_dates.empty: return ['N/A']
        periods = sorted([p for p in valid_dates.dt.to_period(period_freq).unique() if pd.notna(p)])
        return ['All Periods'] + [str(p) for p in periods]
    except Exception as e:
        st.warning(f"Could not determine time periods for {date_column} with frequency {freq}: {e}"); return ['N/A']

def calculate_sla_percentage(df_filtered, target_type, period_col_name):
    if df_filtered is None or df_filtered.empty: return pd.DataFrame({'Time Period': [], f'% {target_type} SLA Met': []})
    if 'Target' not in df_filtered.columns or 'Has breached' not in df_filtered.columns: return pd.DataFrame({'Time Period': [], f'% {target_type} SLA Met': []})
    sla_df = df_filtered[df_filtered['Target'] == target_type].copy()
    if sla_df.empty or sla_df['Has breached'].isnull().all() or period_col_name not in sla_df.columns: return pd.DataFrame({'Time Period': [], f'% {target_type} SLA Met': []})
    count_col = 'Task' if 'Task' in sla_df.columns else 'index'
    if count_col == 'index' and 'index' not in sla_df.columns: sla_df = sla_df.reset_index()
    if count_col not in sla_df.columns: return pd.DataFrame({'Time Period': [], f'% {target_type} SLA Met': []})
    try:
        # True means breached (met is False), False means not breached (met is True)
        sla_df['Has breached_bool'] = sla_df['Has breached'].map({True: False, False: True, pd.NA: False, 'TRUE': False, 'FALSE': True}) 
        sla_grouped = sla_df.groupby(period_col_name).agg(total=(count_col, 'nunique'), met=('Has breached_bool', 'sum')).reset_index()
    except KeyError: return pd.DataFrame({'Time Period': [], f'% {target_type} SLA Met': []})
    sla_grouped[f'% {target_type} SLA Met'] = sla_grouped.apply(lambda row: (row['met'] / row['total'] * 100) if row['total'] > 0 else 0, axis=1)
    sla_grouped = sla_grouped.rename(columns={period_col_name: 'Time Period'})
    return sla_grouped[['Time Period', f'% {target_type} SLA Met']].sort_values('Time Period')

def calculate_kpis(df):
    kpis = {
        "total_accounts": 0,
        "total_opened": 0,
        "total_closed": 0,
        "overall_resolution_sla_met": 0.0,
        "best_resolution_sla_month_value": 0.0,
        "best_resolution_sla_month_name": "N/A",
        "overall_response_sla_met": 0.0, # NEW KPI
        "best_response_sla_month_value": 0.0, # NEW KPI
        "best_response_sla_month_name": "N/A" # NEW KPI
    }
    if df is None or df.empty: return kpis

    count_col = 'Task' if 'Task' in df.columns else 'index'
    if count_col == 'index' and 'index' not in df.columns : df = df.reset_index()

    # Total Unique Accounts
    if 'Account' in df.columns: kpis["total_accounts"] = df['Account'].nunique()

    # Opened Tickets
    if count_col in df.columns: kpis["total_opened"] = df[count_col].nunique()
    else: kpis["total_opened"] = len(df)

    # Closed Tickets
    closed_states_list = ['Closed', 'Resolved', 'Closed Complete']
    if 'State' in df.columns:
        kpis["total_closed"] = df[df['State'].isin(closed_states_list)][count_col].nunique() if count_col in df.columns else len(df[df['State'].isin(closed_states_list)])
    elif 'Closed' in df.columns and pd.api.types.is_datetime64_any_dtype(df['Closed']):
         kpis["total_closed"] = df[df['Closed'].notna()][count_col].nunique() if count_col in df.columns else len(df[df['Closed'].notna()])

    if 'Target' in df.columns and 'Has breached' in df.columns:
        # Overall Resolution SLA Met (%)
        resolution_df = df[(df['Target'] == 'Resolution') & df['Has breached'].notna()].copy()
        if not resolution_df.empty:
            total_resolution = len(resolution_df)
            met_resolution = (resolution_df['Has breached'].map({True: False, False: True, pd.NA: False, 'TRUE': False, 'FALSE': True})).sum()
            kpis["overall_resolution_sla_met"] = (met_resolution / total_resolution * 100) if total_resolution > 0 else 0.0

        # Overall Response SLA Met (%) - NEW
        response_df = df[(df['Target'] == 'Response') & df['Has breached'].notna()].copy()
        if not response_df.empty:
            total_response = len(response_df)
            met_response = (response_df['Has breached'].map({True: False, False: True, pd.NA: False, 'TRUE': False, 'FALSE': True})).sum()
            kpis["overall_response_sla_met"] = (met_response / total_response * 100) if total_response > 0 else 0.0

    if 'Opened_YearMonth' in df.columns and 'Target' in df.columns and 'Has breached' in df.columns and count_col in df.columns:
        # Best Resolution SLA Month
        res_sla_df = df[(df['Target'] == 'Resolution') & df['Has breached'].notna() & df['Opened_YearMonth'].notna()].copy()
        if not res_sla_df.empty and 'Opened_YearMonth' in res_sla_df.columns:
                 res_sla_df['Has breached_bool'] = res_sla_df['Has breached'].map({True: False, False: True, pd.NA: False, 'TRUE': False, 'FALSE': True})
                 monthly_sla = res_sla_df.groupby('Opened_YearMonth').agg(total=(count_col, 'nunique'), met=('Has breached_bool', 'sum')).reset_index()
                 monthly_sla['SLA %'] = monthly_sla.apply(lambda row: (row['met'] / row['total'] * 100) if row['total'] > 0 else 0, axis=1)
                 if not monthly_sla.empty:
                     best_month_data = monthly_sla.loc[monthly_sla['SLA %'].idxmax()]
                     kpis["best_resolution_sla_month_value"] = best_month_data['SLA %']
                     kpis["best_resolution_sla_month_name"] = str(best_month_data['Opened_YearMonth'])

        # Best Response SLA Month - NEW
        resp_sla_df = df[(df['Target'] == 'Response') & df['Has breached'].notna() & df['Opened_YearMonth'].notna()].copy()
        if not resp_sla_df.empty and 'Opened_YearMonth' in resp_sla_df.columns:
                 resp_sla_df['Has breached_bool'] = resp_sla_df['Has breached'].map({True: False, False: True, pd.NA: False, 'TRUE': False, 'FALSE': True})
                 monthly_resp_sla = resp_sla_df.groupby('Opened_YearMonth').agg(total=(count_col, 'nunique'), met=('Has breached_bool', 'sum')).reset_index()
                 monthly_resp_sla['SLA %'] = monthly_resp_sla.apply(lambda row: (row['met'] / row['total'] * 100) if row['total'] > 0 else 0, axis=1)
                 if not monthly_resp_sla.empty:
                     best_month_data_resp = monthly_resp_sla.loc[monthly_resp_sla['SLA %'].idxmax()]
                     kpis["best_response_sla_month_value"] = best_month_data_resp['SLA %']
                     kpis["best_response_sla_month_name"] = str(best_month_data_resp['Opened_YearMonth'])
    return kpis

# --- DOCX Export Function ---
def create_export_docx(kpi_data, figures_data, global_filters_desc, local_filters_desc_map, genai_summaries):
    doc = Document()

    # Define styles if not present (optional, but good for consistency)
    try:
        heading_style_level_5 = doc.styles['Heading 5']
    except KeyError: # Style not found
        heading_style_level_5 = doc.styles.add_style('Heading 5', WD_STYLE_TYPE.PARAGRAPH)
        heading_style_level_5.base_style = doc.styles['Heading 4'] # Example base
        heading_style_level_5.font.size = Pt(10)
        heading_style_level_5.font.bold = True


    # --- Document Title and Header ---
    doc.add_heading('ServiceNow Advanced Analysis Report', level=0)
    p_timestamp = doc.add_paragraph()
    run_timestamp = p_timestamp.add_run(f'Report generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    run_timestamp.font.size = Pt(10)
    run_timestamp.italic = True
    p_timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # --- Global Filters ---
    doc.add_heading('Global Filters Applied:', level=2)
    doc.add_paragraph(global_filters_desc if global_filters_desc else "No global filters applied.")
    doc.add_paragraph()

    # --- KPIs ---
    doc.add_heading('Key Performance Indicators', level=2)
    kpi_list_data = {
        "Total Unique Accounts (in selection)": f"{kpi_data.get('total_accounts', 'N/A'):,}",
        "Opened Tickets (in selection)": f"{kpi_data.get('total_opened', 'N/A'):,}",
        "Closed Tickets (in selection)": f"{kpi_data.get('total_closed', 'N/A'):,}",
        "Overall Resolution SLA Met (%)": f"{kpi_data.get('overall_resolution_sla_met', 0.0):.1f}%",
        f"Best Resolution SLA Month ({kpi_data.get('best_resolution_sla_month_name', 'N/A')})": f"{kpi_data.get('best_resolution_sla_month_value', 0.0):.1f}%",
        "Overall Response SLA Met (%)": f"{kpi_data.get('overall_response_sla_met', 0.0):.1f}%", # NEW
        f"Best Response SLA Month ({kpi_data.get('best_response_sla_month_name', 'N/A')})": f"{kpi_data.get('best_response_sla_month_value', 0.0):.1f}%" # NEW
    }
    for key, value in kpi_list_data.items():
        doc.add_paragraph(f"{key}: {value}", style='ListBullet')
    doc.add_paragraph()

    # --- Figures ---
    for i, item in enumerate(figures_data):
        fig_obj = item['object']
        title = item['title']
        purpose = item.get('purpose', "N/A")
        benefit = item.get('benefit', "N/A")
        req_id = item.get('req_id', f"chart_{i}") # Fallback req_id

        doc.add_heading(title, level=3)

        p_purpose = doc.add_paragraph()
        p_purpose.add_run("Purpose: ").bold = True
        p_purpose.add_run(purpose)

        p_benefit = doc.add_paragraph()
        p_benefit.add_run("Benefit: ").bold = True
        p_benefit.add_run(benefit)

        local_desc = local_filters_desc_map.get(req_id)
        if local_desc:
            p_local_filters = doc.add_paragraph()
            p_local_filters.add_run("Local Filters: ").bold = True
            p_local_filters.add_run(local_desc)
        doc.add_paragraph() # Spacing before image

        if fig_obj is not None:
            try:
                # Using scale as in original PDF export for consistency
                img_bytes = pio.to_image(fig_obj, format='png', scale=1.5) # Consider width/height if scale is problematic
                img_file = io.BytesIO(img_bytes)
                doc.add_picture(img_file, width=Inches(6.2)) # Adjust width as needed for A4/Letter
                img_file.close()
            except Exception as e:
                error_msg = f"Error exporting figure '{title}' to image for Word: {e}"
                doc.add_paragraph(error_msg)
                print(error_msg) # Also print to console for server-side logs if deployed
        else:
            doc.add_paragraph("No data available for this chart with the current filters.")
        doc.add_paragraph() # Spacing after image

        # Add GenAI Summary
        genai_summary = genai_summaries.get(req_id)
        if genai_summary:
            doc.add_heading('GenAI Insights:', level=4)
            summary_lines = genai_summary.split('\n')
            for line_content in summary_lines:
                line = line_content.strip()
                if not line:
                    continue # Skip empty lines

                # Basic Markdown to DOCX conversion (headings and lists)
                if line.startswith("1.  **Chart Overview:**"):
                    doc.add_heading("Chart Overview", level=5) # Use actual heading style
                    line = line.replace("1.  **Chart Overview:**", "").strip()
                elif line.startswith("2.  **Key Insights & Trends:**"):
                    doc.add_heading("Key Insights & Trends", level=5)
                    line = line.replace("2.  **Key Insights & Trends:**", "").strip()
                elif line.startswith("3.  **Potential Implications & Predictive Thoughts:**"):
                    doc.add_heading("Potential Implications & Predictive Thoughts", level=5)
                    line = line.replace("3.  **Potential Implications & Predictive Thoughts:**", "").strip()
                elif line.startswith("4.  **Actionable Recommendations:**"):
                    doc.add_heading("Actionable Recommendations", level=5)
                    line = line.replace("4.  **Actionable Recommendations:**", "").strip()

                if line: # Add remaining content of the line or if it's a normal line
                    p = doc.add_paragraph()
                    if line.startswith("* ") or line.startswith("- "):
                        p.text = line[2:] # Remove bullet character
                        p.style = 'ListBullet'
                    elif line.startswith("**") and line.endswith("**") and len(line) > 4:
                        p.add_run(line[2:-2]).bold = True
                    else:
                        p.add_run(line) # Add as plain text
            doc.add_paragraph() # Spacing after GenAI summary

        doc.add_page_break()

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()

# --- Color Palette ---
PROFESSIONAL_PALETTE = px.colors.qualitative.Plotly
COLOR_RESOLVED = '#1f77b4'
COLOR_SLA_RESPONSE = '#2ca02c' # Green for Response
COLOR_SLA_RESOLUTION = '#ff7f0e' # Orange for Resolution
COLOR_BREACH = px.colors.qualitative.Set1[0] # Red for Breaches 
CLOSED_STATES_LIST = ['Closed', 'Resolved', 'Closed Complete']

# --- New helper to clear all GenAI summaries in session state ---
def clear_genai_summaries_session_state():
    for key in list(st.session_state.keys()):
        if key.startswith("genai_summary_"):
            del st.session_state[key]
    if 'all_genai_summaries_for_export' in st.session_state:
        del st.session_state['all_genai_summaries_for_export']

# --- Helper function for dynamic period formatting in charts (especially 6.1) ---
def format_period_label(date_obj, freq):
    """Formats a datetime object into a string label based on the given frequency."""
    if freq == 'M':
        return date_obj.strftime('%Y-%m')
    elif freq == 'W':
        # ISO week and year: YYYY-Www (e.g., 2023-W01)
        # Ensure week number is zero-padded
        return date_obj.strftime('%Y-W%U') # %U for Sunday as first day of week
    elif freq == 'Q':
        # Quarter format: YYYY-Qx (e.g., 2023-Q1)
        return f"{date_obj.year}-Q{date_obj.quarter}"
    elif freq == 'Y':
        return date_obj.strftime('%Y')
    else:
        return date_obj.strftime('%Y-%m-%d') # Default fallback


# --- Main Application ---
st.title("📊 ServiceNow CSM Operational KPI Dashboard")
st.markdown("Upload your ServiceNow data export (Excel format) to visualize key metrics and trends.")

uploaded_file = st.file_uploader("Choose an Excel file (.xlsx)", type="xlsx")

if 'data' not in st.session_state: st.session_state.data = None
if 'uploaded_file_name' not in st.session_state: st.session_state.uploaded_file_name = None

if uploaded_file is not None:
    if st.session_state.uploaded_file_name != uploaded_file.name:
        with st.spinner("Loading and processing data..."):
            st.session_state.data = load_data(uploaded_file)
            st.session_state.uploaded_file_name = uploaded_file.name
            keys_to_clear_on_new_file = ['global_periods', 'global_regions', 'global_accounts', 'current_freq']
            clear_genai_summaries_session_state() # Clear all GenAI summaries
            for key in keys_to_clear_on_new_file: st.session_state.pop(key, None)

df_original = st.session_state.data

if df_original is not None and not df_original.empty:
    # --- Global Sidebar Filters ---
    st.sidebar.header("Global Filters")
    global_time_col = 'Opened'
    if global_time_col not in df_original.columns:
        st.sidebar.error(f"'{global_time_col}' column missing, cannot apply time filters."); st.stop()

    time_slice_options_global = {'Month': 'M', 'Week': 'W', 'Quarter': 'Q', 'Year': 'Y'}
    selected_time_slice_global_label = st.sidebar.selectbox(
        "Time Slice (based on 'Opened' date):", options=list(time_slice_options_global.keys()), index=0, key='global_timeframe',
        on_change=clear_genai_summaries_session_state # Call helper function
    )
    freq_global = time_slice_options_global[selected_time_slice_global_label]

    if 'global_periods' not in st.session_state or st.session_state.get('current_freq') != freq_global:
        st.session_state.global_periods = get_time_periods(df_original, global_time_col, freq_global)
        st.session_state.current_freq = freq_global
    selected_period_global = st.sidebar.selectbox(
        f"Specific {selected_time_slice_global_label}:", options=st.session_state.global_periods, index=0, key='global_period',
        on_change=clear_genai_summaries_session_state # Call helper function
    )
    selected_regions_global = [] # Use a list for multi-select
    if 'Region' in df_original.columns:
        if 'global_regions' not in st.session_state:
             # For multiselect, we don't need 'All Regions' in the options list
             st.session_state.global_regions = sorted(df_original['Region'].dropna().unique().tolist())
        selected_regions_global = st.sidebar.multiselect(
            "Regions:", options=st.session_state.global_regions, default=[], key='global_region_multi',
            on_change=clear_genai_summaries_session_state # Call helper function
        )
    else: st.sidebar.info("'Region' column not found. Region filter disabled.")
    selected_accounts_global = []
    if 'Account' in df_original.columns:
        if 'global_accounts' not in st.session_state:
            st.session_state.global_accounts = sorted(df_original['Account'].dropna().unique().tolist())
        selected_accounts_global = st.sidebar.multiselect(
            "Accounts:", options=st.session_state.global_accounts, default=[], key='global_accounts_select',
            on_change=clear_genai_summaries_session_state # Call helper function
        )
    else: st.sidebar.info("'Account' column not found. Account filter disabled.")

    # --- Apply Global Filters ---
    df_filtered_global = df_original.copy()
    global_filters_applied_list = []
    if selected_accounts_global and 'Account' in df_filtered_global.columns:
        df_filtered_global = df_filtered_global[df_filtered_global['Account'].isin(selected_accounts_global)]
        global_filters_applied_list.append(f"Accounts: {', '.join(selected_accounts_global)}")
    elif 'Account' in df_filtered_global.columns : global_filters_applied_list.append("Accounts: All")
    else: global_filters_applied_list.append("Accounts: N/A (column missing)")
    if selected_regions_global and 'Region' in df_filtered_global.columns:
        df_filtered_global = df_filtered_global[df_filtered_global['Region'].isin(selected_regions_global)]
        global_filters_applied_list.append(f"Regions: {', '.join(selected_regions_global)}")
    elif 'Region' in df_filtered_global.columns: global_filters_applied_list.append("Regions: All")
    else: global_filters_applied_list.append("Regions: N/A (column missing)")
    if selected_period_global != 'N/A':
        period_freq_global = freq_global if freq_global != 'Y' else 'A-DEC'
        time_col_present = global_time_col in df_filtered_global.columns and not df_filtered_global[global_time_col].isnull().all()
        if time_col_present:
            df_filtered_global['Global_Time_Period'] = pd.to_datetime(df_filtered_global[global_time_col], errors='coerce', dayfirst=(global_time_col in DAY_FIRST_TRUE_COLS)).dt.to_period(period_freq_global).astype(str)
            if selected_period_global != 'All Periods':
                if 'Global_Time_Period' in df_filtered_global.columns:
                     df_filtered_global = df_filtered_global[df_filtered_global['Global_Time_Period'] == selected_period_global]
            global_filters_applied_list.append(f"Time Slice: {selected_time_slice_global_label}")
            global_filters_applied_list.append(f"Period: {selected_period_global}")
        else:
             global_filters_applied_list.append("Time Slice/Period: N/A (Missing/Empty Date)")
             if 'Global_Time_Period' not in df_filtered_global.columns: df_filtered_global['Global_Time_Period'] = 'N/A'
    else:
        global_filters_applied_list.append("Time Slice/Period: N/A")
        if 'Global_Time_Period' not in df_filtered_global.columns: df_filtered_global['Global_Time_Period'] = 'N/A'
    global_filters_description = ", ".join(filter(None, global_filters_applied_list))
    st.sidebar.caption(f"Current Global Filters: {global_filters_description}")

    figures_for_export = []
    local_filters_desc_map = {}

    # --- Reusable function for chart-specific filters ---
    def apply_local_filters(df_base, chart_id_prefix, primary_date_col, allow_time_slice=True, allow_period=True, allow_region=True, allow_account=True, allow_priority=False, allow_state=False, allow_category=False, allow_subcategory=False):
        df_local = df_base.copy()
        local_filters_applied_list = []
        
        # Determine unique options for dropdowns based on current df_local
        all_regions = sorted(df_local['Region'].dropna().unique().tolist()) if 'Region' in df_local.columns else []
        all_accounts = sorted(df_local['Account'].dropna().unique().tolist()) if 'Account' in df_local.columns else []
        all_priorities = sorted(df_local['Priority'].dropna().unique().tolist()) if 'Priority' in df_local.columns else []
        all_states = sorted(df_local['State'].dropna().unique().tolist()) if 'State' in df_local.columns else []
        all_categories = sorted(df_local['Category'].dropna().unique().tolist()) if 'Category' in df_local.columns else []
        all_subcategories = sorted(df_local['Subcategory'].dropna().unique().tolist()) if 'Subcategory' in df_local.columns else []

        # Determine number of columns for filter layout
        num_filters_enabled = sum([allow_time_slice, allow_period, allow_region, allow_account, allow_priority, allow_state, allow_category, allow_subcategory])
        
        # Use a list to store the columns generated by st.columns, so we can access them by index
        if num_filters_enabled > 0:
            cols_main = st.columns(min(3, num_filters_enabled))
            col_idx = 0
        else:
            cols_main = []
            col_idx = 0

        # Time Slice
        selected_time_slice = None
        current_freq = None
        if allow_time_slice and primary_date_col in df_local.columns:
            if col_idx >= len(cols_main):
                # Add more columns if needed, ensuring we don't exceed 3 columns per row for filters
                cols_main.extend(st.columns(min(3, num_filters_enabled - col_idx))) 
            with cols_main[col_idx]:
                selected_time_slice = st.selectbox("Time Slice:", list(time_slice_options_global.keys()), index=0, key=f'{chart_id_prefix}_ts', on_change=clear_genai_summaries_session_state)
            current_freq = time_slice_options_global[selected_time_slice]
            local_filters_applied_list.append(f"Slice: {selected_time_slice}")
            col_idx += 1
        elif allow_time_slice:
            local_filters_applied_list.append("Slice: N/A (Date col missing)")

        # Period
        selected_period = 'All Periods'
        period_col_name = None
        if allow_period and primary_date_col in df_local.columns and selected_time_slice:
            if col_idx >= len(cols_main):
                cols_main.extend(st.columns(min(3, num_filters_enabled - col_idx)))
            with cols_main[col_idx]:
                periods = get_time_periods(df_local, primary_date_col, current_freq)
                selected_period = st.selectbox(f"{selected_time_slice} (Local):", periods, index=0, key=f'{chart_id_prefix}_p', on_change=clear_genai_summaries_session_state, disabled=(not periods or periods == ['N/A']))
            period_freq = current_freq if current_freq != 'Y' else 'A-DEC'
            period_col_name = f'Local_Time_Period_{chart_id_prefix}'
            
            if not df_local[primary_date_col].isnull().all():
                df_local[period_col_name] = pd.to_datetime(df_local[primary_date_col], errors='coerce', dayfirst=(primary_date_col in DAY_FIRST_TRUE_COLS)).dt.to_period(period_freq).astype(str)
                if selected_period != 'All Periods' and selected_period != 'N/A':
                    df_local = df_local[df_local[period_col_name] == selected_period]
            else:
                df_local[period_col_name] = 'N/A' # Mark all as N/A if date column is empty
            local_filters_applied_list.append(f"Period: {selected_period}")
            col_idx += 1
        elif allow_period:
            local_filters_applied_list.append("Period: N/A (Date col missing)")

        # Region
                # Region (Multi-select)
        selected_regions = []
        if allow_region and 'Region' in df_local.columns:
            if col_idx >= len(cols_main):
                cols_main.extend(st.columns(min(3, num_filters_enabled - col_idx)))
            with cols_main[col_idx]:
                # We don't need 'All Regions' in the multiselect list itself
                selected_regions = st.multiselect("Regions (Local):", all_regions, default=[], key=f'{chart_id_prefix}_reg_multi', on_change=clear_genai_summaries_session_state, disabled=(not all_regions))
            if selected_regions: # An empty list means 'All'
                df_local = df_local[df_local['Region'].isin(selected_regions)]
                local_filters_applied_list.append(f"Regions: {', '.join(selected_regions)}")
            else:
                local_filters_applied_list.append("Regions: All")
            col_idx += 1
        elif allow_region:
            local_filters_applied_list.append("Regions: N/A (column missing)")

        # Account
        selected_accounts = []
        if allow_account and 'Account' in df_local.columns:
            if col_idx >= len(cols_main):
                cols_main.extend(st.columns(min(3, num_filters_enabled - col_idx)))
            with cols_main[col_idx]:
                selected_accounts = st.multiselect("Accounts (Local):", all_accounts, default=[], key=f'{chart_id_prefix}_acc', on_change=clear_genai_summaries_session_state, disabled=(not all_accounts))
            if selected_accounts:
                df_local = df_local[df_local['Account'].isin(selected_accounts)]
                local_filters_applied_list.append(f"Accounts: {', '.join(selected_accounts)}")
            else:
                local_filters_applied_list.append("Accounts: All")
            col_idx += 1
        elif allow_account:
            local_filters_applied_list.append("Accounts: N/A (column missing)")

        # Priority
        selected_priority = 'All Priorities'
        if allow_priority and 'Priority' in df_local.columns:
            if col_idx >= len(cols_main):
                cols_main.extend(st.columns(min(3, num_filters_enabled - col_idx)))
            with cols_main[col_idx]:
                selected_priority = st.selectbox("Priority (Local):", ['All Priorities'] + all_priorities, index=0, key=f'{chart_id_prefix}_prio', on_change=clear_genai_summaries_session_state, disabled=(not all_priorities))
            if selected_priority != 'All Priorities':
                df_local = df_local[df_local['Priority'] == selected_priority]
            local_filters_applied_list.append(f"Priority: {selected_priority}")
            col_idx += 1
        elif allow_priority:
            local_filters_applied_list.append("Priority: N/A (column missing)")

        # State
        selected_state = 'All States'
        if allow_state and 'State' in df_local.columns:
            if col_idx >= len(cols_main):
                cols_main.extend(st.columns(min(3, num_filters_enabled - col_idx)))
            with cols_main[col_idx]:
                selected_state = st.selectbox("State (Local):", ['All States'] + all_states, index=0, key=f'{chart_id_prefix}_state', on_change=clear_genai_summaries_session_state, disabled=(not all_states))
            if selected_state != 'All States':
                df_local = df_local[df_local['State'] == selected_state]
            local_filters_applied_list.append(f"State: {selected_state}")
            col_idx += 1
        elif allow_state:
            local_filters_applied_list.append("State: N/A (column missing)")
            
        # Category
        selected_category = 'All Categories'
        if allow_category and 'Category' in df_local.columns:
            if col_idx >= len(cols_main):
                cols_main.extend(st.columns(min(3, num_filters_enabled - col_idx)))
            with cols_main[col_idx]:
                selected_category = st.selectbox("Category (Local):", ['All Categories'] + all_categories, index=0, key=f'{chart_id_prefix}_cat', on_change=clear_genai_summaries_session_state, disabled=(not all_categories))
            if selected_category != 'All Categories':
                df_local = df_local[df_local['Category'] == selected_category]
            local_filters_applied_list.append(f"Category: {selected_category}")
            col_idx += 1
        elif allow_category:
            local_filters_applied_list.append("Category: N/A (column missing)")

        # Subcategory
        selected_subcategory = 'All Subcategories'
        if allow_subcategory and 'Subcategory' in df_local.columns:
            # Subcategory options should react to selected_category if Category filter is available
            filtered_subcats = all_subcategories
            if allow_category and selected_category != 'All Categories' and 'Category' in df_local.columns:
                filtered_subcats = sorted(df_local[df_local['Category']==selected_category]['Subcategory'].dropna().unique().tolist())

            if col_idx >= len(cols_main):
                cols_main.extend(st.columns(min(3, num_filters_enabled - col_idx)))
            with cols_main[col_idx]:
                selected_subcategory = st.selectbox("Subcategory (Local):", ['All Subcategories'] + filtered_subcats, index=0, key=f'{chart_id_prefix}_subcat', on_change=clear_genai_summaries_session_state, disabled=(not filtered_subcats))
            if selected_subcategory != 'All Subcategories':
                df_local = df_local[df_local['Subcategory'] == selected_subcategory]
            local_filters_applied_list.append(f"Subcategory: {selected_subcategory}")
            col_idx += 1
        elif allow_subcategory:
            local_filters_applied_list.append("Subcategory: N/A (column missing)")

        local_filters_description = ", ".join(filter(None, local_filters_applied_list))
        st.caption(f"Local Filters: {local_filters_description}")
        return df_local, local_filters_description, period_col_name, selected_time_slice, current_freq

    # --- Function to display chart and GenAI button (modified) ---
    def display_chart_with_genai(chart_obj, chart_df, chart_title, chart_purpose, chart_benefit, chart_id, chart_type="general", x_col=None, y_col=None, value_col=None, name_col=None):
        if chart_obj:
            st.plotly_chart(chart_obj, use_container_width=True)
            genai_button_key = f"genai_btn_{chart_id}"
            genai_summary_key = f"genai_summary_{chart_id}"
            
            # Display summary directly if generated or generate on button click
            if st.button(f"🤖 Generate GenAI Insights for '{chart_title}'", key=genai_button_key):
                with st.spinner(f"Generating AI insights for '{chart_title}'..."):
                    data_summary_str = generate_data_summary_for_llm(chart_df, chart_type, x_col, y_col, value_col, name_col)
                    summary = get_genai_summary(chart_title, chart_purpose, chart_benefit, data_summary_str, chart_type)
                    st.session_state[genai_summary_key] = summary
                    # Also store in the aggregated summaries for export
                    if 'all_genai_summaries_for_export' not in st.session_state:
                         st.session_state['all_genai_summaries_for_export'] = {}
                    st.session_state['all_genai_summaries_for_export'][chart_id] = summary
            
            if genai_summary_key in st.session_state:
                st.markdown(st.session_state[genai_summary_key])
        else:
            st.info(f"No data available to display or generate insights for '{chart_title}' with current filters.")

    # --- NEW: Function to generate all GenAI summaries for export ---
    def generate_all_chart_genai_summaries(figures_list_param):
        # Initialize a placeholder for GenAI summaries in session state if not present
        if 'all_genai_summaries_for_export' not in st.session_state:
            st.session_state['all_genai_summaries_for_export'] = {}

        charts_to_process_for_genai = []
        for chart_info in figures_list_param:
            chart_id = chart_info.get("req_id")
            if chart_id:
                # If summary already exists (either from individual button or previous batch run), reuse it
                if chart_id not in st.session_state['all_genai_summaries_for_export']:
                    # Check if individual button click already stored it
                    if f"genai_summary_{chart_id}" in st.session_state:
                        st.session_state['all_genai_summaries_for_export'][chart_id] = st.session_state[f"genai_summary_{chart_id}"]
                    else:
                        charts_to_process_for_genai.append(chart_info)
        
        if not charts_to_process_for_genai:
            st.info("All chart GenAI summaries already generated or no new charts to process.")
            return st.session_state['all_genai_summaries_for_export']

        progress_bar = st.progress(0, text="Generating GenAI insights for charts...")
        charts_processed_count = 0

        for i, chart_info in enumerate(charts_to_process_for_genai):
            chart_id = chart_info.get("req_id")
            chart_title = chart_info.get("title")
            chart_purpose = chart_info.get("purpose")
            chart_benefit = chart_info.get("benefit")
            chart_df = chart_info.get("chart_df")
            chart_type = chart_info.get("chart_type")
            x_col = chart_info.get("x_col")
            y_col = chart_info.get("y_col")
            value_col = chart_info.get("value_col")
            name_col = chart_info.get("name_col")

            progress_text = f"Generating AI insights for '{chart_title}' ({charts_processed_count + 1}/{len(charts_to_process_for_genai)})..."
            progress_bar.progress((charts_processed_count + 1) / len(charts_to_process_for_genai), text=progress_text)
            
            if chart_info.get("object") is None: # If chart could not be generated due to no data
                summary = f"No chart generated for '{chart_title}' due to missing or filtered data, so no GenAI summary could be produced."
            else:
                data_summary_str = generate_data_summary_for_llm(chart_df, chart_type, x_col, y_col, value_col, name_col)
                summary = get_genai_summary(chart_title, chart_purpose, chart_benefit, data_summary_str, chart_type)
            
            st.session_state['all_genai_summaries_for_export'][chart_id] = summary
            charts_processed_count += 1
            # A short sleep might be beneficial for API rate limits and also for user experience
            # Groq's client handles rate limiting, but a small delay is fine.
            # import time
            # time.sleep(0.1) # Example: sleep for 100ms between calls

        progress_bar.empty() # Hide the progress bar
        return st.session_state['all_genai_summaries_for_export']


    # --- ALL YOUR CHART SECTIONS GO HERE ---

    # --- 📊 Overall Performance & KPIs ---
    st.header("📊 Overall Performance & KPIs")
    st.caption(f"Key metrics reflecting overall service performance based on global filters: {global_filters_description}")
    kpi_data = calculate_kpis(df_filtered_global)
    
    # Updated to 7 columns for new KPIs
    kpi_col1, kpi_col2, kpi_col3, kpi_col4, kpi_col5, kpi_col6, kpi_col7 = st.columns(7) 
    
    kpi_col1.metric(label="Total Unique Accounts", value=f"{kpi_data['total_accounts']:,}")
    kpi_col2.metric(label="Opened Tickets", value=f"{kpi_data['total_opened']:,}")
    kpi_col3.metric(label="Closed Tickets", value=f"{kpi_data['total_closed']:,}")
    kpi_col4.metric(label="Overall Resolution SLA Met (%)", value=f"{kpi_data['overall_resolution_sla_met']:.1f}%")
    kpi_col5.metric(label=f"Best Resolution SLA Month ({kpi_data['best_resolution_sla_month_name']})", value=f"{kpi_data['best_resolution_sla_month_value']:.1f}%")
    
    # NEW KPI Cards
    kpi_col6.metric(label="Overall Response SLA Met (%)", value=f"{kpi_data['overall_response_sla_met']:.1f}%")
    kpi_col7.metric(label=f"Best Response SLA Month ({kpi_data['best_response_sla_month_name']})", value=f"{kpi_data['best_response_sla_month_value']:.1f}%")
    st.markdown("---")

    # --- 🌍 Account & Regional Insights ---
    st.header("🌍 Account & Regional Insights")

    st.subheader("1.1 Managed Service Accounts Overview")
    purpose_req1 = "To show the distribution of unique customer accounts across different geographical regions."
    benefit_req1 = "Helps understand regional customer footprint and concentration, supporting targeted regional strategies or resource allocation."
    st.caption(f"**Purpose:** {purpose_req1} / **Benefit:** {benefit_req1}")
    
    # Local filters for req1
    df_r1_local, lfilt_r1_desc, _, _, _ = apply_local_filters(
        df_filtered_global, 'req1', 'Opened', 
        allow_time_slice=True, allow_period=True, allow_region=True, allow_account=True,
        allow_priority=False, allow_state=False, allow_category=False, allow_subcategory=False 
    )
    local_filters_desc_map['req1'] = lfilt_r1_desc

    col1_r1, col2_r1 = st.columns([1,2])
    
    r1_title_chart = "Accounts per Region"
    chart_data_r1_df = None 

    with col1_r1:
        st.markdown("###### Unique Accounts by Region Table")
        if 'Account' in df_r1_local.columns and 'Region' in df_r1_local.columns:
            acc_df_r1_table = df_r1_local[['Region','Account']].drop_duplicates().dropna()
            if not acc_df_r1_table.empty:
                reg_counts_r1 = acc_df_r1_table.groupby('Region',as_index=False)['Account'].nunique().rename(columns={'Account':'Unique Accounts'})
                total_row_r1 = pd.DataFrame([{'Region':'TOTAL (All Regions)','Unique Accounts':acc_df_r1_table['Account'].nunique()}])
                st.dataframe(pd.concat([reg_counts_r1.sort_values('Region'),total_row_r1],ignore_index=True),hide_index=True,use_container_width=True)
                st.markdown("###### Explore Accounts in a Region")
                regions_list_r1_drill = ['All'] + sorted(acc_df_r1_table['Region'].unique().tolist())
                drill_reg_r1 = st.selectbox("Select Region for drill-down:",regions_list_r1_drill,key='r1_drill_region')
                accounts_to_show = acc_df_r1_table[acc_df_r1_table['Region']==drill_reg_r1]['Account'].unique() if drill_reg_r1!='All' else acc_df_r1_table['Account'].unique()
                st.dataframe(pd.DataFrame({'Account Name':sorted(accounts_to_show)}),height=200,hide_index=True,use_container_width=True)
            else: st.info("No account data for table.")
        else: st.warning("Missing 'Account' or 'Region' for table.")
    with col2_r1:
        st.markdown(f"###### Chart: {r1_title_chart}")
        fig1 = None
        if 'Account' in df_r1_local.columns and 'Region' in df_r1_local.columns:
            chart_data_r1_df = df_r1_local.groupby('Region',as_index=False)['Account'].nunique().rename(columns={'Account':'Unique Accounts', 'Region':'Region'})
            if not chart_data_r1_df.empty:
                fig1 = px.bar(chart_data_r1_df,x='Region',y='Unique Accounts',title=r1_title_chart,text_auto=True,color='Region',color_discrete_sequence=PROFESSIONAL_PALETTE)
                fig1.update_layout(showlegend=False,xaxis={'categoryorder':'total descending'}).update_traces(textposition='outside')
                display_chart_with_genai(fig1, chart_data_r1_df, r1_title_chart, purpose_req1, benefit_req1, "req1", chart_type="bar", x_col="Region", y_col="Unique Accounts")
            else: st.info("No account data to plot (filtered).")
        else: st.info("Chart needs 'Account' & 'Region' (filtered).")
    figures_for_export.append({
        "title": "1.1 Managed Service Accounts by Region",
        "object": fig1,
        "req_id": "req1",
        "purpose": purpose_req1,
        "benefit": benefit_req1,
        "chart_df": chart_data_r1_df,
        "chart_type": "bar",
        "x_col": "Region",
        "y_col": "Unique Accounts",
        "value_col": None,
        "name_col": None
    })
    st.markdown("---")

    # --- 📦 Ticket Volume & Workload ---
    st.header("📦 Ticket Volume & Workload")
    st.subheader("2.1 Opened vs Closed Tickets Trend")
    purpose_req2 = "To compare the number of tickets opened versus tickets closed over selected time periods."
    benefit_req2 = "Helps identify trends in workload and operational efficiency."
    st.caption(f"**Purpose:** {purpose_req2} / **Benefit:** {benefit_req2}")
    
    # Local filters for req2
    df_r2_local, lfilt_r2_desc, period_col_r2, selected_time_slice_r2, freq_r2 = apply_local_filters(
        df_filtered_global, 'req2', 'Opened', 
        allow_time_slice=True, allow_period=True, allow_region=True, allow_account=True,
        allow_priority=True, allow_state=True, allow_category=True, allow_subcategory=True
    )
    local_filters_desc_map['req2'] = lfilt_r2_desc

    fig2 = None; compare_df_r2 = None
    r2_title = f"Opened vs Closed Tickets ({selected_time_slice_r2})"
    st.markdown(f"###### Chart: {r2_title}")
    
    count_col_r2='Task' if 'Task' in df_r2_local.columns else 'index'
    if count_col_r2=='index' and 'index' not in df_r2_local.columns: df_r2_local=df_r2_local.reset_index()
    
    if df_r2_local.empty: st.info("No data for local filters (Opened vs Closed).")
    else:
        # Re-apply period filtering logic for opened and closed, as `apply_local_filters` handles primary_date_col ('Opened')
        opened_df = df_r2_local.groupby(period_col_r2,as_index=False)[count_col_r2].nunique().rename(columns={period_col_r2:'Time Period',count_col_r2:'Opened Count'})
        
        closed_df_final = pd.DataFrame({'Time Period':[],'Closed Count':[]})
        if 'Closed' in df_r2_local.columns and pd.api.types.is_datetime64_any_dtype(df_r2_local['Closed']):
            p_freq_r2 = freq_r2 if freq_r2 != 'Y' else 'A-DEC'
            df_r2_local['Period_Closed_r2_chart'] = pd.to_datetime(df_r2_local['Closed'], errors='coerce', dayfirst=('Closed' in DAY_FIRST_TRUE_COLS)).dt.to_period(p_freq_r2).astype(str)
            closed_data = df_r2_local[df_r2_local['Period_Closed_r2_chart']!='N/A']
            closed_df_final = closed_data.groupby('Period_Closed_r2_chart',as_index=False)[count_col_r2].nunique().rename(columns={'Period_Closed_r2_chart':'Time Period',count_col_r2:'Closed Count'})
        elif 'State' in df_r2_local.columns:
            closed_state_data = df_r2_local[df_r2_local['State'].isin(CLOSED_STATES_LIST) & (df_r2_local[period_col_r2]!='N/A')]
            if not closed_state_data.empty: closed_df_final = closed_state_data.groupby(period_col_r2,as_index=False)[count_col_r2].nunique().rename(columns={period_col_r2:'Time Period',count_col_r2:'Closed Count'})
        
        compare_df_r2 = pd.merge(opened_df,closed_df_final,on='Time Period',how='outer').fillna(0).sort_values('Time Period')
        if compare_df_r2.empty: st.info("No Opened/Closed data for local filters.")
        else:
            fig2 = px.bar(compare_df_r2,x='Time Period',y=['Opened Count','Closed Count'],title=r2_title,barmode='group',color_discrete_sequence=PROFESSIONAL_PALETTE)
            fig2.update_layout(yaxis_title="Number of Tickets",legend_title_text='Status',hovermode="x unified")
            display_chart_with_genai(fig2, compare_df_r2, r2_title, purpose_req2, benefit_req2, "req2", chart_type="bar_grouped", x_col="Time Period", y_col=['Opened Count','Closed Count'])
    figures_for_export.append({
        "title": "2.1 Opened vs Closed Tickets Trend",
        "object": fig2,
        "req_id": "req2",
        "purpose": purpose_req2,
        "benefit": benefit_req2,
        "chart_df": compare_df_r2,
        "chart_type": "bar_grouped",
        "x_col": "Time Period",
        "y_col": ['Opened Count','Closed Count'],
        "value_col": None,
        "name_col": None
    })
    st.markdown("---")

    st.subheader("2.2 Ticket Volume by Account")
    p_add5, b_add5 = "Identify accounts generating most tickets.", "Understand key customer engagement, focus support."
    st.caption(f"**Purpose:** {p_add5} / **Benefit:** {b_add5}")

    # Local filters for add5
    df_add5_local, lfilt_add5_desc, _, _, _ = apply_local_filters(
        df_filtered_global, 'add5', 'Opened',
        allow_time_slice=True, allow_period=True, allow_region=True, allow_account=False, # Account is the X-axis
        allow_priority=True, allow_state=True, allow_category=True, allow_subcategory=True
    )
    local_filters_desc_map['add5'] = lfilt_add5_desc
    
    add_fig5 = None; acc_counts_add5 = None; top_n_acc_val = 15 # Default value
    if 'Account' in df_add5_local.columns and 'Task' in df_add5_local.columns:
        top_n_acc_val = st.slider("Top N Accounts:",1,30,15,key="add5_slider")
        acc_counts_add5 = df_add5_local.groupby('Account')['Task'].nunique().nlargest(top_n_acc_val).reset_index(name='Ticket Count')
        if not acc_counts_add5.empty:
            add_fig5_title = f"Ticket Volume by Account (Top {top_n_acc_val})"
            add_fig5 = px.bar(acc_counts_add5,x='Account',y='Ticket Count',title=add_fig5_title,text_auto=True)
            add_fig5.update_layout(xaxis={'categoryorder':'total descending'},yaxis_title="Number of Tickets")
            display_chart_with_genai(add_fig5, acc_counts_add5, add_fig5_title, p_add5, b_add5, "add5", chart_type="bar", x_col="Account", y_col="Ticket Count")
        else: st.info("No account data (filtered).")
    else: st.info("Missing 'Account' or 'Task'.")
    figures_for_export.append({
        "title": f"2.2 Ticket Volume by Account (Top {top_n_acc_val})",
        "object": add_fig5,
        "req_id": "add5",
        "purpose": p_add5,
        "benefit": b_add5,
        "chart_df": acc_counts_add5,
        "chart_type": "bar",
        "x_col": "Account",
        "y_col": "Ticket Count",
        "value_col": None,
        "name_col": None
    })
    st.markdown("---")

    st.subheader("2.3 Distribution of Currently Open Tickets by State")
    p_add6, b_add6 = "Show breakdown of open tickets by current status.", "Insight into operational pipeline, identify bottlenecks."
    st.caption(f"**Purpose:** {p_add6} / **Benefit:** {b_add6}")

    # Local filters for add6
    df_add6_local, lfilt_add6_desc, _, _, _ = apply_local_filters(
        df_filtered_global, 'add6', 'Opened',
        allow_time_slice=True, allow_period=True, allow_region=True, allow_account=True,
        allow_priority=True, allow_state=False, allow_category=True, allow_subcategory=True # State is the X-axis
    )
    local_filters_desc_map['add6'] = lfilt_add6_desc

    add_fig6 = None; state_counts_add6 = None
    add_fig6_title = "Distribution of Open Tickets by State"
    if 'State' in df_add6_local.columns and 'Task' in df_add6_local.columns:
        open_t_add6 = df_add6_local[~df_add6_local['State'].isin(CLOSED_STATES_LIST)].copy()
        if not open_t_add6.empty:
            state_counts_add6 = open_t_add6.groupby('State')['Task'].nunique().reset_index(name='Count')
            if not state_counts_add6.empty:
                add_fig6 = px.pie(state_counts_add6,names='State',values='Count',title=add_fig6_title,hole=0.3)
                add_fig6.update_traces(textposition='inside',textinfo='percent+label')
                display_chart_with_genai(add_fig6, state_counts_add6, add_fig6_title, p_add6, b_add6, "add6", chart_type="pie", name_col="State", value_col="Count")
            else: st.info("No open tickets by state (filtered).")
        else: st.info("No open tickets (filtered).")
    else: st.info("Missing 'State' or 'Task'.")
    figures_for_export.append({
        "title": add_fig6_title,
        "object": add_fig6,
        "req_id": "add6",
        "purpose": p_add6,
        "benefit": b_add6,
        "chart_df": state_counts_add6,
        "chart_type": "pie",
        "x_col": None,
        "y_col": None,
        "value_col": "Count",
        "name_col": "State"
    })
    st.markdown("---")

    st.subheader("2.4 Ticket Distribution by Priority")
    p_add7, b_add7 = "Understand proportion of tickets at each priority.", "Gauge workload urgency, align resources."
    st.caption(f"**Purpose:** {p_add7} / **Benefit:** {b_add7}")

    # Local filters for add7
    df_add7_local, lfilt_add7_desc, _, _, _ = apply_local_filters(
        df_filtered_global, 'add7', 'Opened',
        allow_time_slice=True, allow_period=True, allow_region=True, allow_account=True,
        allow_priority=False, allow_state=True, allow_category=True, allow_subcategory=True # Priority is the X-axis
    )
    local_filters_desc_map['add7'] = lfilt_add7_desc

    add_fig7 = None; prio_counts_add7 = None
    add_fig7_title = "Ticket Distribution by Priority"
    if 'Priority' in df_add7_local.columns and 'Task' in df_add7_local.columns:
        prio_counts_add7 = df_add7_local.groupby('Priority')['Task'].nunique().reset_index(name='Count')
        if not prio_counts_add7.empty:
            try: prio_counts_add7['SortKey']=prio_counts_add7['Priority'].astype(str).str.extract(r'(\d+)').astype(int); prio_counts_add7=prio_counts_add7.sort_values('SortKey')
            except: prio_counts_add7=prio_counts_add7.sort_values('Priority') 
            add_fig7 = px.pie(prio_counts_add7,names='Priority',values='Count',title=add_fig7_title,hole=0.3)
            add_fig7.update_traces(textposition='inside',textinfo='percent+label')
            display_chart_with_genai(add_fig7, prio_counts_add7, add_fig7_title, p_add7, b_add7, "add7", chart_type="pie", name_col="Priority", value_col="Count")
        else: st.info("No priority data (filtered).")
    else: st.info("Missing 'Priority' or 'Task'.")
    figures_for_export.append({
        "title": add_fig7_title,
        "object": add_fig7,
        "req_id": "add7",
        "purpose": p_add7,
        "benefit": b_add7,
        "chart_df": prio_counts_add7,
        "chart_type": "pie",
        "x_col": None,
        "y_col": None,
        "value_col": "Count",
        "name_col": "Priority"
    })
    st.markdown("---")

    st.subheader("2.5 Open Ticket Aging Distribution")
    p_add8, b_add8 = "Categorize open tickets by age.", "Highlight aging tickets needing attention."
    st.caption(f"**Purpose:** {p_add8} / **Benefit:** {b_add8}")

    # Local filters for add8
    df_add8_local, lfilt_add8_desc, _, _, _ = apply_local_filters(
        df_filtered_global, 'add8', 'Opened',
        allow_time_slice=True, allow_period=True, allow_region=True, allow_account=True,
        allow_priority=True, allow_state=False, allow_category=True, allow_subcategory=True # State handled by current open tickets logic
    )
    local_filters_desc_map['add8'] = lfilt_add8_desc

    add_fig8 = None; aging_counts_df = None
    add_fig8_title = "Open Ticket Aging Distribution"
    if 'State' in df_add8_local.columns and 'Opened' in df_add8_local.columns and 'Task' in df_add8_local.columns:
        open_aging_df = df_add8_local[~df_add8_local['State'].isin(CLOSED_STATES_LIST)].copy()
        if not open_aging_df.empty:
            open_aging_df['Opened'] = pd.to_datetime(open_aging_df['Opened'],errors='coerce',dayfirst=('Opened' in DAY_FIRST_TRUE_COLS))
            open_aging_df = open_aging_df.dropna(subset=['Opened'])
            if not open_aging_df.empty:
                now_utc = datetime.now(timezone.utc)
                open_aging_df['Opened_utc'] = open_aging_df['Opened'].dt.tz_localize('UTC',ambiguous='NaT',nonexistent='NaT') if open_aging_df['Opened'].dt.tz is None else open_aging_df['Opened'].dt.tz_convert('UTC')
                open_aging_df = open_aging_df.dropna(subset=['Opened_utc'])
                if not open_aging_df.empty:
                    open_aging_df['Age (Days)'] = (now_utc - open_aging_df['Opened_utc']).dt.total_seconds()/(60*60*24)
                    open_aging_df = open_aging_df[open_aging_df['Age (Days)']>=0]
                    bins = [0,1,3,7,14,30,np.inf]; labels = ['0-1 Day','1-3 Days','3-7 Days','7-14 Days','14-30 Days','30+ Days']
                    open_aging_df['Age Bin'] = pd.cut(open_aging_df['Age (Days)'],bins=bins,labels=labels,right=False)
                    aging_counts_df = open_aging_df.groupby('Age Bin',observed=False)['Task'].nunique().reset_index(name='Count')
                    if not aging_counts_df.empty:
                        add_fig8 = px.bar(aging_counts_df,x='Age Bin',y='Count',title=add_fig8_title,text_auto=True)
                        add_fig8.update_layout(xaxis_title="Ticket Age",yaxis_title="Number of Open Tickets")
                        display_chart_with_genai(add_fig8, aging_counts_df, add_fig8_title, p_add8, b_add8, "add8", chart_type="bar", x_col="Age Bin", y_col="Count")
                    else: st.info("Could not calculate aging (filtered).")
                else: st.info("No open tickets with valid UTC 'Opened' dates.")
            else: st.info("No open tickets with valid 'Opened' dates.")
        else: st.info("No open tickets found (filtered).")
    else: st.info("Missing 'State', 'Opened', or 'Task'.")
    figures_for_export.append({
        "title": add_fig8_title,
        "object": add_fig8,
        "req_id": "add8",
        "purpose": p_add8,
        "benefit": b_add8,
        "chart_df": aging_counts_df,
        "chart_type": "bar",
        "x_col": "Age Bin",
        "y_col": "Count",
        "value_col": None,
        "name_col": None
    })
    st.markdown("---")
        # --- NEW: Advanced Open Ticket Aging Analysis ---
    st.subheader("2.5.1 Advanced Open Ticket Aging Analysis")
    st.caption("The following charts break down the open ticket aging distribution by different dimensions like Region, Priority, and Account to identify specific areas with aging backlogs.")

    # We can reuse the same local filter set for all advanced aging charts.
    # This ensures consistency for this sub-section.
    df_add8_adv_local, lfilt_add8_adv_desc, _, _, _ = apply_local_filters(
        df_filtered_global, 'add8_adv', 'Opened',
        allow_time_slice=True, allow_period=True, allow_region=True, allow_account=True, # Note: region filter is now multi-select
        allow_priority=True, allow_state=False, allow_category=True, allow_subcategory=True
    )
    local_filters_desc_map['add8_adv_common'] = lfilt_add8_adv_desc # Common filter description for this section

    # --- Core Aging Calculation Logic (reused for all sub-charts) ---
    open_aging_adv_df = None
    if 'State' in df_add8_adv_local.columns and 'Opened' in df_add8_adv_local.columns and 'Task' in df_add8_adv_local.columns:
        temp_df = df_add8_adv_local[~df_add8_adv_local['State'].isin(CLOSED_STATES_LIST)].copy()
        if not temp_df.empty:
            temp_df['Opened'] = pd.to_datetime(temp_df['Opened'],errors='coerce',dayfirst=('Opened' in DAY_FIRST_TRUE_COLS))
            temp_df = temp_df.dropna(subset=['Opened'])
            if not temp_df.empty:
                now_utc = datetime.now(timezone.utc)
                temp_df['Opened_utc'] = temp_df['Opened'].dt.tz_localize('UTC',ambiguous='NaT',nonexistent='NaT') if temp_df['Opened'].dt.tz is None else temp_df['Opened'].dt.tz_convert('UTC')
                temp_df = temp_df.dropna(subset=['Opened_utc'])
                if not temp_df.empty:
                    temp_df['Age (Days)'] = (now_utc - temp_df['Opened_utc']).dt.total_seconds()/(60*60*24)
                    temp_df = temp_df[temp_df['Age (Days)']>=0]
                    bins = [0,1,3,7,14,30,np.inf]; labels = ['0-1 Day','1-3 Days','3-7 Days','7-14 Days','14-30 Days','30+ Days']
                    temp_df['Age Bin'] = pd.cut(temp_df['Age (Days)'],bins=bins,labels=labels,right=False)
                    open_aging_adv_df = temp_df.copy() # This is the base df for our new charts

    # --- Chart: Aging by Region ---
    st.markdown("###### Open Ticket Aging by Region")
    p_add8_reg, b_add8_reg = "Break down open ticket aging by geographical region.", "Highlights which regions are struggling with older tickets, informing targeted interventions."
    st.caption(f"**Purpose:** {p_add8_reg} / **Benefit:** {b_add8_reg}")

    fig_add8_region = None
    aging_by_region_df = None
    fig_add8_region_title = "Open Ticket Aging by Region"

    if open_aging_adv_df is not None and 'Region' in open_aging_adv_df.columns:
        aging_by_region_df = open_aging_adv_df.groupby(['Age Bin', 'Region'], observed=False)['Task'].nunique().reset_index(name='Count')
        if not aging_by_region_df.empty:
            fig_add8_region = px.bar(aging_by_region_df, x='Age Bin', y='Count', color='Region',
                                     title=fig_add8_region_title, barmode='stack',
                                     labels={'Count': 'Number of Open Tickets', 'Age Bin': 'Ticket Age'},
                                     color_discrete_sequence=px.colors.qualitative.Vivid)
            fig_add8_region.update_layout(xaxis_title="Ticket Age", yaxis_title="Number of Open Tickets")
            display_chart_with_genai(fig_add8_region, aging_by_region_df, fig_add8_region_title, p_add8_reg, b_add8_reg, "add8_region", chart_type="bar", x_col="Age Bin", y_col="Count")
        else:
            st.info("No data to display for aging by region with current filters.")
    else:
        st.info("Required data for 'Aging by Region' (e.g., 'Region', 'State', 'Opened') is missing or filtered out.")

    figures_for_export.append({
        "title": fig_add8_region_title, "object": fig_add8_region, "req_id": "add8_region",
        "purpose": p_add8_reg, "benefit": b_add8_reg, "chart_df": aging_by_region_df,
        "chart_type": "bar", "x_col": "Age Bin", "y_col": "Count", "name_col": "Region"
    })
    st.markdown("---")

    # --- Chart: Aging by Priority ---
    st.markdown("###### Open Ticket Aging by Priority")
    p_add8_prio, b_add8_prio = "Analyze the age of open tickets based on their priority level.", "Identifies if high-priority tickets are being left unresolved, signaling potential process failures."
    st.caption(f"**Purpose:** {p_add8_prio} / **Benefit:** {b_add8_prio}")

    fig_add8_prio = None
    aging_by_prio_df = None
    fig_add8_prio_title = "Open Ticket Aging by Priority"

    if open_aging_adv_df is not None and 'Priority' in open_aging_adv_df.columns:
        aging_by_prio_df = open_aging_adv_df.groupby(['Age Bin', 'Priority'], observed=False)['Task'].nunique().reset_index(name='Count')
        if not aging_by_prio_df.empty:
            # Sort priorities logically
            try:
                aging_by_prio_df['SortKey']=aging_by_prio_df['Priority'].astype(str).str.extract(r'(\d+)').astype(int)
                priority_order = aging_by_prio_df.sort_values('SortKey')['Priority'].unique()
            except:
                priority_order = sorted(aging_by_prio_df['Priority'].unique())

            fig_add8_prio = px.bar(aging_by_prio_df, x='Age Bin', y='Count', color='Priority',
                                   title=fig_add8_prio_title, barmode='stack',
                                   category_orders={"Priority": priority_order},
                                   labels={'Count': 'Number of Open Tickets', 'Age Bin': 'Ticket Age'},
                                   color_discrete_sequence=px.colors.qualitative.Bold)
            fig_add8_prio.update_layout(xaxis_title="Ticket Age", yaxis_title="Number of Open Tickets")
            display_chart_with_genai(fig_add8_prio, aging_by_prio_df, fig_add8_prio_title, p_add8_prio, b_add8_prio, "add8_prio", chart_type="bar", x_col="Age Bin", y_col="Count")
        else:
            st.info("No data to display for aging by priority with current filters.")
    else:
        st.info("Required data for 'Aging by Priority' (e.g., 'Priority', 'State', 'Opened') is missing or filtered out.")

    figures_for_export.append({
        "title": fig_add8_prio_title, "object": fig_add8_prio, "req_id": "add8_prio",
        "purpose": p_add8_prio, "benefit": b_add8_prio, "chart_df": aging_by_prio_df,
        "chart_type": "bar", "x_col": "Age Bin", "y_col": "Count", "name_col": "Priority"
    })
    st.markdown("---")

    # --- Chart: Aging by Top N Accounts ---
    st.markdown("###### Open Ticket Aging by Top N Accounts")
    p_add8_acc, b_add8_acc = "Visualize aging tickets for the most active accounts.", "Focuses attention on key customers who may be experiencing service degradation due to old tickets."
    st.caption(f"**Purpose:** {p_add8_acc} / **Benefit:** {b_add8_acc}")

    fig_add8_acc = None
    aging_by_acc_df = None
    fig_add8_acc_title = "Open Ticket Aging by Top N Accounts" # Default

    if open_aging_adv_df is not None and 'Account' in open_aging_adv_df.columns:
        # Determine Top N accounts based on the number of OPEN tickets
        top_n_accounts = st.slider("Select Top N Accounts to Display:", min_value=3, max_value=20, value=10, key="add8_top_n_slider")
        top_accounts_list = open_aging_adv_df['Account'].value_counts().nlargest(top_n_accounts).index.tolist()
        
        # Filter the aging dataframe for only these top accounts
        filtered_aging_for_accounts = open_aging_adv_df[open_aging_adv_df['Account'].isin(top_accounts_list)]
        
        aging_by_acc_df = filtered_aging_for_accounts.groupby(['Age Bin', 'Account'], observed=False)['Task'].nunique().reset_index(name='Count')

        if not aging_by_acc_df.empty:
            fig_add8_acc_title = f"Open Ticket Aging by Top {top_n_accounts} Accounts"
            fig_add8_acc = px.bar(aging_by_acc_df, x='Age Bin', y='Count', color='Account',
                                  title=fig_add8_acc_title, barmode='stack',
                                  labels={'Count': 'Number of Open Tickets', 'Age Bin': 'Ticket Age'},
                                  color_discrete_sequence=px.colors.qualitative.Alphabet)
            fig_add8_acc.update_layout(xaxis_title="Ticket Age", yaxis_title="Number of Open Tickets")
            display_chart_with_genai(fig_add8_acc, aging_by_acc_df, fig_add8_acc_title, p_add8_acc, b_add8_acc, "add8_account", chart_type="bar", x_col="Age Bin", y_col="Count")
        else:
            st.info("No data to display for aging by account with current filters.")
    else:
        st.info("Required data for 'Aging by Account' (e.g., 'Account', 'State', 'Opened') is missing or filtered out.")

    figures_for_export.append({
        "title": fig_add8_acc_title, "object": fig_add8_acc, "req_id": "add8_account",
        "purpose": p_add8_acc, "benefit": b_add8_acc, "chart_df": aging_by_acc_df,
        "chart_type": "bar", "x_col": "Age Bin", "y_col": "Count", "name_col": "Account"
    })
    st.markdown("---")

    st.subheader("2.6 Tickets by Assignment Group")
    p_new1, b_new1 = "Visualize ticket distribution across assignment groups.", "Understand workload balance, identify overburdened teams."
    st.caption(f"**Purpose:** {p_new1} / **Benefit:** {b_new1}")

    # Local filters for new1
    df_new1_local, lfilt_new1_desc, _, _, _ = apply_local_filters(
        df_filtered_global, 'new1', 'Opened',
        allow_time_slice=True, allow_period=True, allow_region=True, allow_account=True,
        allow_priority=True, allow_state=True, allow_category=True, allow_subcategory=True
    )
    local_filters_desc_map['new1'] = lfilt_new1_desc

    fig_new1 = None; assign_counts_new1_df = None
    fig_new1_title = "Ticket Count by Assignment Group" # Default title
    if 'Assignment group' in df_new1_local.columns and 'Task' in df_new1_local.columns:
        assign_counts_new1 = df_new1_local.groupby('Assignment group')['Task'].nunique().reset_index(name='Ticket Count').sort_values('Ticket Count',ascending=False)
        if not assign_counts_new1.empty:
            num_grps = len(assign_counts_new1); disp_grps_df = assign_counts_new1.head(15) if num_grps>15 else assign_counts_new1
            fig_new1_title = "Ticket Count by Assignment Group" + (" (Top 15)" if num_grps>15 else "")
            fig_new1 = px.bar(disp_grps_df,x='Assignment group',y='Ticket Count',title=fig_new1_title,text_auto=True)
            fig_new1.update_layout(xaxis_title="Assignment Group",yaxis_title="Number of Tickets",xaxis={'categoryorder':'total descending'})
            display_chart_with_genai(fig_new1, disp_grps_df, fig_new1_title, p_new1, b_new1, "new1", chart_type="bar", x_col="Assignment group", y_col="Ticket Count")
            assign_counts_new1_df = disp_grps_df # Use this as the chart_df for export
        else: st.info("No ticket data by assignment group (filtered).")
    else: st.info("Missing 'Assignment group' or 'Task'.")
    figures_for_export.append({
        "title": fig_new1_title,
        "object": fig_new1,
        "req_id": "new1",
        "purpose": p_new1,
        "benefit": b_new1,
        "chart_df": assign_counts_new1_df, 
        "chart_type": "bar",
        "x_col": "Assignment group",
        "y_col": "Ticket Count",
        "value_col": None,
        "name_col": None
    })
    st.markdown("---")

    # --- ⏱️ SLA Performance & Compliance ---
    st.header("⏱️ SLA Performance & Compliance")

    # --- NEW RESPONSE CHARTS ---
    st.subheader("3.1 'Response' Target Tickets Opened Volume")
    p_new_resp_vol, b_new_resp_vol = "Track volume of 'Response' SLA tickets opened over selected time periods.", "Indicates throughput for acknowledgment-bound tickets."
    st.caption(f"**Purpose:** {p_new_resp_vol} / **Benefit:** {b_new_resp_vol}")

    df_new_resp_vol_local, lfilt_new_resp_vol_desc, period_col_new_resp_vol, selected_time_slice_new_resp_vol, freq_new_resp_vol = apply_local_filters(
        df_filtered_global, 'new_resp_vol', 'Opened', # Using 'Opened' date for Response
        allow_time_slice=True, allow_period=True, allow_region=True, allow_account=True,
        allow_priority=True, allow_state=True, allow_category=True, allow_subcategory=True
    )
    local_filters_desc_map['new_resp_vol'] = lfilt_new_resp_vol_desc

    fig_new_resp_vol = None; resp_counts_new_resp_vol_df = None
    title_new_resp_vol = f"'Response' Target Tickets Opened ({selected_time_slice_new_resp_vol})"
    st.markdown(f"###### Chart: {title_new_resp_vol}")
    
    count_col_new_resp_vol='Task' if 'Task' in df_new_resp_vol_local.columns else 'index'
    if count_col_new_resp_vol=='index' and 'index' not in df_new_resp_vol_local.columns: df_new_resp_vol_local=df_new_resp_vol_local.reset_index()

    if df_new_resp_vol_local.empty: st.info("No data for local filters (Response Target Opened).")
    elif not all(c in df_new_resp_vol_local for c in['Target',count_col_new_resp_vol,period_col_new_resp_vol]) or df_new_resp_vol_local[period_col_new_resp_vol].eq('N/A').all(): st.warning("Missing columns/periods (Response Target Opened).")
    else:
        resp_counts_new_resp_vol_df = df_new_resp_vol_local[(df_new_resp_vol_local['Target']=='Response') & (df_new_resp_vol_local[period_col_new_resp_vol]!='N/A')].groupby(period_col_new_resp_vol,as_index=False)[count_col_new_resp_vol].nunique().rename(columns={count_col_new_resp_vol:'Tickets Opened',period_col_new_resp_vol:'Time Period'}).sort_values('Time Period')
        if resp_counts_new_resp_vol_df.empty: st.info("No 'Response' tickets found opened (local filters).")
        else:
            fig_new_resp_vol=px.bar(resp_counts_new_resp_vol_df,x='Time Period',y='Tickets Opened',title=title_new_resp_vol,text_auto=True)
            fig_new_resp_vol.update_layout(yaxis_title="Tickets Opened").update_traces(marker_color=COLOR_SLA_RESPONSE,textposition='outside')
            display_chart_with_genai(fig_new_resp_vol, resp_counts_new_resp_vol_df, title_new_resp_vol, p_new_resp_vol, b_new_resp_vol, "new_resp_vol", chart_type="bar", x_col="Time Period", y_col="Tickets Opened")
    figures_for_export.append({
        "title":"3.1 'Response' Target Tickets Opened Volume",
        "object":fig_new_resp_vol,
        "req_id":"new_resp_vol",
        "purpose":p_new_resp_vol,
        "benefit":b_new_resp_vol,
        "chart_df": resp_counts_new_resp_vol_df,
        "chart_type": "bar",
        "x_col": "Time Period",
        "y_col": "Tickets Opened",
        "value_col": None,
        "name_col": None
    })
    st.markdown("---")

    st.subheader("3.2 SLA Response Performance (%)")
    p_r4, b_r4 = "Measure % of tickets meeting 'Response' SLA.", "Indicates effectiveness in acknowledging tickets."
    st.caption(f"**Purpose:** {p_r4} / **Benefit:** {b_r4}")

    # Local filters for req4 (now req_sla_resp_perc)
    df_r4_local, lfilt_r4_desc, period_col_r4, selected_time_slice_r4, freq_r4 = apply_local_filters(
        df_filtered_global, 'req_sla_resp_perc', 'Opened', # Note: Using 'Opened' date for SLA calculation
        allow_time_slice=True, allow_period=True, allow_region=True, allow_account=True,
        allow_priority=True, allow_state=True, allow_category=True, allow_subcategory=True
    )
    local_filters_desc_map['req_sla_resp_perc'] = lfilt_r4_desc

    fig4 = None; resp_sla_perc_df = None
    r4_title=f"% Response SLA Met ({selected_time_slice_r4})"
    st.markdown(f"###### Chart: {r4_title}")

    if df_r4_local.empty or period_col_r4 is None or df_r4_local[period_col_r4].eq('N/A').all(): st.info("No data for local filters (Response SLA).")
    elif not all(c in df_r4_local for c in ['Target','Has breached',period_col_r4]): st.warning("Missing columns/periods (Response SLA).")
    else:
        resp_sla_perc_df=calculate_sla_percentage(df_r4_local,'Response',period_col_r4)
        if resp_sla_perc_df.empty or resp_sla_perc_df.iloc[:,1].sum()==0: st.info("Not enough data for Response SLA %.")
        else:
            fig4=px.bar(resp_sla_perc_df,x='Time Period',y='% Response SLA Met',title=r4_title,text='% Response SLA Met')
            fig4.update_layout(yaxis_title="% Met",yaxis_range=[0,105]).update_traces(texttemplate='%{text:.1f}%',textposition='outside',marker_color=COLOR_SLA_RESPONSE)
            display_chart_with_genai(fig4, resp_sla_perc_df, r4_title, p_r4, b_r4, "req_sla_resp_perc", chart_type="bar", x_col="Time Period", y_col="% Response SLA Met")
    figures_for_export.append({
        "title":"3.2 SLA Response Performance (%)",
        "object":fig4,
        "req_id":"req_sla_resp_perc",
        "purpose":p_r4,
        "benefit":b_r4,
        "chart_df": resp_sla_perc_df,
        "chart_type": "bar",
        "x_col": "Time Period",
        "y_col": "% Response SLA Met",
        "value_col": None,
        "name_col": None
    })
    st.markdown("---")

    st.subheader("3.3 Response SLA Breach Count by Region")
    p_new_resp_brch_reg, b_new_resp_brch_reg = "Identify number of 'Response' SLA breaches per region.", "Pinpoint regions with higher acknowledgment SLA failure rates."
    st.caption(f"**Purpose:** {p_new_resp_brch_reg} / **Benefit:** {b_new_resp_brch_reg}")

    # Local filters for new_resp_brch_reg
    df_new_resp_brch_reg_local, lfilt_new_resp_brch_reg_desc, _, _, _ = apply_local_filters(
        df_filtered_global, 'new_resp_brch_reg', 'Opened',
        allow_time_slice=True, allow_period=True, allow_region=False, allow_account=True, # Region is the X-axis
        allow_priority=True, allow_state=True, allow_category=True, allow_subcategory=True
    )
    local_filters_desc_map['new_resp_brch_reg'] = lfilt_new_resp_brch_reg_desc

    fig_new_resp_brch_reg = None; breach_counts_new_resp_brch_reg_df = None
    title_new_resp_brch_reg = "Response SLA Breach Count by Region"
    if 'Has breached' in df_new_resp_brch_reg_local.columns and 'Region' in df_new_resp_brch_reg_local.columns and 'Target' in df_new_resp_brch_reg_local.columns:
        sla_df_new_resp_brch_reg = df_new_resp_brch_reg_local[df_new_resp_brch_reg_local['Target']=='Response'].copy()
        if pd.api.types.is_bool_dtype(sla_df_new_resp_brch_reg['Has breached']) or pd.api.types.is_object_dtype(sla_df_new_resp_brch_reg['Has breached']): 
            sla_df_new_resp_brch_reg['Breached_Flag'] = sla_df_new_resp_brch_reg['Has breached'].map({True:True,pd.NA:False,False:False, 'TRUE':True, 'FALSE':False}).astype(bool)
            breach_counts_new_resp_brch_reg_df = sla_df_new_resp_brch_reg[sla_df_new_resp_brch_reg['Breached_Flag']].groupby('Region').size().reset_index(name='Breach Count')
            if not breach_counts_new_resp_brch_reg_df.empty:
                fig_new_resp_brch_reg=px.bar(breach_counts_new_resp_brch_reg_df,x='Region',y='Breach Count',title=title_new_resp_brch_reg,text_auto=True)
                fig_new_resp_brch_reg.update_layout(xaxis={'categoryorder':'total descending'},yaxis_title="Number of Breaches").update_traces(marker_color=COLOR_BREACH)
                display_chart_with_genai(fig_new_resp_brch_reg, breach_counts_new_resp_brch_reg_df, title_new_resp_brch_reg, p_new_resp_brch_reg, b_new_resp_brch_reg, "new_resp_brch_reg", chart_type="bar", x_col="Region", y_col="Breach Count")
            else: st.info("No Response SLA breaches (filtered).")
        else: st.warning("'Has breached' not boolean-like for Breach Count.")
    else: st.info("Missing 'Has breached', 'Region', or 'Target'.")
    figures_for_export.append({
        "title":title_new_resp_brch_reg,
        "object":fig_new_resp_brch_reg,
        "req_id":"new_resp_brch_reg",
        "purpose":p_new_resp_brch_reg,
        "benefit":b_new_resp_brch_reg,
        "chart_df": breach_counts_new_resp_brch_reg_df,
        "chart_type": "bar",
        "x_col": "Region",
        "y_col": "Breach Count",
        "value_col": None,
        "name_col": None
    })
    st.markdown("---")

    st.subheader("3.4 Response SLA Breach Rate by Priority")
    p_new_resp_brch_prio, b_new_resp_brch_prio = "Identify if priorities are prone to breaching Response SLAs.", "Pinpoint priorities needing process improvements for acknowledgment."
    st.caption(f"**Purpose:** {p_new_resp_brch_prio} / **Benefit:** {b_new_resp_brch_prio}")

    # Local filters for new_resp_brch_prio
    df_new_resp_brch_prio_local, lfilt_new_resp_brch_prio_desc, _, _, _ = apply_local_filters(
        df_filtered_global, 'new_resp_brch_prio', 'Opened',
        allow_time_slice=True, allow_period=True, allow_region=True, allow_account=True,
        allow_priority=False, allow_state=True, allow_category=True, allow_subcategory=True # Priority is the X-axis
    )
    local_filters_desc_map['new_resp_brch_prio'] = lfilt_new_resp_brch_prio_desc

    fig_new_resp_brch_prio = None; prio_summary_new_resp_brch_prio_df = None
    title_new_resp_brch_prio = "Response SLA Breach Rate by Priority"
    if 'Priority' in df_new_resp_brch_prio_local.columns and 'Target' in df_new_resp_brch_prio_local.columns and 'Has breached' in df_new_resp_brch_prio_local.columns and 'Task' in df_new_resp_brch_prio_local.columns:
        sla_prio_df_new_resp_brch_prio = df_new_resp_brch_prio_local[df_new_resp_brch_prio_local['Target']=='Response'].copy()
        if not sla_prio_df_new_resp_brch_prio.empty:
            sla_prio_df_new_resp_brch_prio['Is_Breached'] = sla_prio_df_new_resp_brch_prio['Has breached'].map({True:1,False:0,pd.NA:0, 'TRUE':1, 'FALSE':0}).astype(int)
            prio_summary_new_resp_brch_prio_df = sla_prio_df_new_resp_brch_prio.groupby('Priority').agg(Total_SLAs=('Task','nunique'),Breached_SLAs=('Is_Breached','sum')).reset_index()
            prio_summary_new_resp_brch_prio_df['Breach Rate (%)'] = prio_summary_new_resp_brch_prio_df.apply(lambda r: (r['Breached_SLAs']/r['Total_SLAs']*100) if r['Total_SLAs']>0 else 0, axis=1)
            try: prio_summary_new_resp_brch_prio_df['SortKey']=prio_summary_new_resp_brch_prio_df['Priority'].astype(str).str.extract(r'(\d+)').astype(int); prio_summary_new_resp_brch_prio_df=prio_summary_new_resp_brch_prio_df.sort_values('SortKey')
            except: prio_summary_new_resp_brch_prio_df=prio_summary_new_resp_brch_prio_df.sort_values('Priority')
            if not prio_summary_new_resp_brch_prio_df.empty:
                fig_new_resp_brch_prio=px.bar(prio_summary_new_resp_brch_prio_df,x='Priority',y='Breach Rate (%)',title=title_new_resp_brch_prio)
                fig_new_resp_brch_prio.update_layout(yaxis_title="Breach Rate (%)",yaxis_range=[0,105]).update_traces(texttemplate='%{y:.1f}%', textposition='outside', marker_color=COLOR_BREACH)
                display_chart_with_genai(fig_new_resp_brch_prio, prio_summary_new_resp_brch_prio_df, title_new_resp_brch_prio, p_new_resp_brch_prio, b_new_resp_brch_prio, "new_resp_brch_prio", chart_type="bar", x_col="Priority", y_col="Breach Rate (%)")
            else: st.info("No data for SLA breach rate by priority (filtered).")
        else: st.info("No 'Response' target tickets for SLA breach rate by priority.")
    else: st.info("Missing 'Priority', 'Target', 'Has breached', or 'Task'.")
    figures_for_export.append({
        "title":title_new_resp_brch_prio,
        "object":fig_new_resp_brch_prio,
        "req_id":"new_resp_brch_prio",
        "purpose":p_new_resp_brch_prio,
        "benefit":b_new_resp_brch_prio,
        "chart_df": prio_summary_new_resp_brch_prio_df,
        "chart_type": "bar",
        "x_col": "Priority",
        "y_col": "Breach Rate (%)",
        "value_col": None,
        "name_col": None
    })
    st.markdown("---")

    # --- RESOLUTION CHARTS (RENUMBERED) ---
    st.subheader("3.5 'Resolution' Target Tickets Closed Volume")
    p_req_sla_res_vol, b_req_sla_res_vol = "Track volume of 'Resolution' SLA tickets closed over selected time periods.", "Indicates throughput for SLA-bound tickets."
    st.caption(f"**Purpose:** {p_req_sla_res_vol} / **Benefit:** {b_req_sla_res_vol}")

    # Local filters for req_sla_res_vol (formerly req3)
    df_req_sla_res_vol_local, lfilt_req_sla_res_vol_desc, period_col_req_sla_res_vol, selected_time_slice_req_sla_res_vol, freq_req_sla_res_vol = apply_local_filters(
        df_filtered_global, 'req_sla_res_vol', 'Closed', # Note: Using 'Closed' date for this chart
        allow_time_slice=True, allow_period=True, allow_region=True, allow_account=True,
        allow_priority=True, allow_state=True, allow_category=True, allow_subcategory=True
    )
    local_filters_desc_map['req_sla_res_vol'] = lfilt_req_sla_res_vol_desc

    fig_req_sla_res_vol = None; res_counts_req_sla_res_vol_df = None
    title_req_sla_res_vol=f"'Resolution' Target Tickets Closed ({selected_time_slice_req_sla_res_vol})"
    st.markdown(f"###### Chart: {title_req_sla_res_vol}")
    
    count_col_req_sla_res_vol='Task' if 'Task' in df_req_sla_res_vol_local.columns else 'index'
    if count_col_req_sla_res_vol=='index' and 'index' not in df_req_sla_res_vol_local.columns: df_req_sla_res_vol_local=df_req_sla_res_vol_local.reset_index()

    if df_req_sla_res_vol_local.empty: st.info("No data for local filters (Res. Target Closed).")
    elif not all(c in df_req_sla_res_vol_local for c in['Target',count_col_req_sla_res_vol,period_col_req_sla_res_vol]) or df_req_sla_res_vol_local[period_col_req_sla_res_vol].eq('N/A').all(): st.warning("Missing columns/periods (Res. Target Closed).")
    else:
        res_counts_req_sla_res_vol_df = df_req_sla_res_vol_local[(df_req_sla_res_vol_local['Target']=='Resolution') & (df_req_sla_res_vol_local[period_col_req_sla_res_vol]!='N/A')].groupby(period_col_req_sla_res_vol,as_index=False)[count_col_req_sla_res_vol].nunique().rename(columns={count_col_req_sla_res_vol:'Tickets Closed',period_col_req_sla_res_vol:'Time Period'}).sort_values('Time Period')
        if res_counts_req_sla_res_vol_df.empty: st.info("No 'Resolution' tickets found closed (local filters).")
        else:
            fig_req_sla_res_vol=px.bar(res_counts_req_sla_res_vol_df,x='Time Period',y='Tickets Closed',title=title_req_sla_res_vol,text_auto=True)
            fig_req_sla_res_vol.update_layout(yaxis_title="Tickets Closed").update_traces(marker_color=COLOR_RESOLVED,textposition='outside')
            display_chart_with_genai(fig_req_sla_res_vol, res_counts_req_sla_res_vol_df, title_req_sla_res_vol, p_req_sla_res_vol, b_req_sla_res_vol, "req_sla_res_vol", chart_type="bar", x_col="Time Period", y_col="Tickets Closed")
    figures_for_export.append({
        "title":"3.5 'Resolution' Target Tickets Closed Volume",
        "object":fig_req_sla_res_vol,
        "req_id":"req_sla_res_vol",
        "purpose":p_req_sla_res_vol,
        "benefit":b_req_sla_res_vol,
        "chart_df": res_counts_req_sla_res_vol_df,
        "chart_type": "bar",
        "x_col": "Time Period",
        "y_col": "Tickets Closed",
        "value_col": None,
        "name_col": None
    })
    st.markdown("---")

    st.subheader("3.6 SLA Resolution Performance (%)")
    p_req_sla_res_perc, b_req_sla_res_perc = "Measure % of tickets meeting 'Resolution' SLA.", "Indicates effectiveness in resolving tickets."
    st.caption(f"**Purpose:** {p_req_sla_res_perc} / **Benefit:** {b_req_sla_res_perc}")

    # Local filters for req_sla_res_perc (formerly req5)
    df_req_sla_res_perc_local, lfilt_req_sla_res_perc_desc, period_col_req_sla_res_perc, selected_time_slice_req_sla_res_perc, freq_req_sla_res_perc = apply_local_filters(
        df_filtered_global, 'req_sla_res_perc', 'Opened', # Note: Using 'Opened' date for SLA calculation
        allow_time_slice=True, allow_period=True, allow_region=True, allow_account=True,
        allow_priority=True, allow_state=True, allow_category=True, allow_subcategory=True
    )
    local_filters_desc_map['req_sla_res_perc'] = lfilt_req_sla_res_perc_desc

    fig_req_sla_res_perc = None; res_sla_perc_df = None
    title_req_sla_res_perc=f"% Resolution SLA Met ({selected_time_slice_req_sla_res_perc})"
    st.markdown(f"###### Chart: {title_req_sla_res_perc}")

    if df_req_sla_res_perc_local.empty or period_col_req_sla_res_perc is None or df_req_sla_res_perc_local[period_col_req_sla_res_perc].eq('N/A').all(): st.info("No data for local filters (Resolution SLA).")
    elif not all(c in df_req_sla_res_perc_local for c in ['Target','Has breached',period_col_req_sla_res_perc]): st.warning("Missing columns/periods (Resolution SLA).")
    else:
        res_sla_perc_df=calculate_sla_percentage(df_req_sla_res_perc_local,'Resolution',period_col_req_sla_res_perc)
        if res_sla_perc_df.empty or res_sla_perc_df.iloc[:,1].sum()==0: st.info("Not enough data for Resolution SLA %.")
        else:
            fig_req_sla_res_perc=px.bar(res_sla_perc_df,x='Time Period',y='% Resolution SLA Met',title=title_req_sla_res_perc,text='% Resolution SLA Met')
            fig_req_sla_res_perc.update_layout(yaxis_title="% Met",yaxis_range=[0,105]).update_traces(texttemplate='%{text:.1f}%',textposition='outside',marker_color=COLOR_SLA_RESOLUTION)
            display_chart_with_genai(fig_req_sla_res_perc, res_sla_perc_df, title_req_sla_res_perc, p_req_sla_res_perc, b_req_sla_res_perc, "req_sla_res_perc", chart_type="bar", x_col="Time Period", y_col="% Resolution SLA Met")
    figures_for_export.append({
        "title":"3.6 SLA Resolution Performance (%)",
        "object":fig_req_sla_res_perc,
        "req_id":"req_sla_res_perc",
        "purpose":p_req_sla_res_perc,
        "benefit":b_req_sla_res_perc,
        "chart_df": res_sla_perc_df,
        "chart_type": "bar",
        "x_col": "Time Period",
        "y_col": "% Resolution SLA Met",
        "value_col": None,
        "name_col": None
    })
    st.markdown("---")

    st.subheader("3.7 Resolution SLA Breach Count by Region")
    p_req_sla_res_brch_reg, b_req_sla_res_brch_reg = "Identify number of 'Resolution' SLA breaches per region.", "Pinpoint regions with higher SLA failure rates."
    st.caption(f"**Purpose:** {p_req_sla_res_brch_reg} / **Benefit:** {b_req_sla_res_brch_reg}")

    # Local filters for req_sla_res_brch_reg (formerly add3)
    df_req_sla_res_brch_reg_local, lfilt_req_sla_res_brch_reg_desc, _, _, _ = apply_local_filters(
        df_filtered_global, 'req_sla_res_brch_reg', 'Opened',
        allow_time_slice=True, allow_period=True, allow_region=False, allow_account=True, # Region is the X-axis
        allow_priority=True, allow_state=True, allow_category=True, allow_subcategory=True
    )
    local_filters_desc_map['req_sla_res_brch_reg'] = lfilt_req_sla_res_brch_reg_desc

    fig_req_sla_res_brch_reg = None; breach_counts_req_sla_res_brch_reg_df = None
    title_req_sla_res_brch_reg = "Resolution SLA Breach Count by Region"
    if 'Has breached' in df_req_sla_res_brch_reg_local.columns and 'Region' in df_req_sla_res_brch_reg_local.columns and 'Target' in df_req_sla_res_brch_reg_local.columns:
        sla_df_req_sla_res_brch_reg = df_req_sla_res_brch_reg_local[df_req_sla_res_brch_reg_local['Target']=='Resolution'].copy()
        if pd.api.types.is_bool_dtype(sla_df_req_sla_res_brch_reg['Has breached']) or pd.api.types.is_object_dtype(sla_df_req_sla_res_brch_reg['Has breached']): 
            sla_df_req_sla_res_brch_reg['Breached_Flag'] = sla_df_req_sla_res_brch_reg['Has breached'].map({True:True,pd.NA:False,False:False, 'TRUE':True, 'FALSE':False}).astype(bool)
            breach_counts_req_sla_res_brch_reg_df = sla_df_req_sla_res_brch_reg[sla_df_req_sla_res_brch_reg['Breached_Flag']].groupby('Region').size().reset_index(name='Breach Count')
            if not breach_counts_req_sla_res_brch_reg_df.empty:
                fig_req_sla_res_brch_reg=px.bar(breach_counts_req_sla_res_brch_reg_df,x='Region',y='Breach Count',title=title_req_sla_res_brch_reg,text_auto=True)
                fig_req_sla_res_brch_reg.update_layout(xaxis={'categoryorder':'total descending'},yaxis_title="Number of Breaches").update_traces(marker_color=COLOR_BREACH)
                display_chart_with_genai(fig_req_sla_res_brch_reg, breach_counts_req_sla_res_brch_reg_df, title_req_sla_res_brch_reg, p_req_sla_res_brch_reg, b_req_sla_res_brch_reg, "req_sla_res_brch_reg", chart_type="bar", x_col="Region", y_col="Breach Count")
            else: st.info("No Res. SLA breaches (filtered).")
        else: st.warning("'Has breached' not boolean-like for Breach Count.")
    else: st.info("Missing 'Has breached', 'Region', or 'Target'.")
    figures_for_export.append({
        "title":title_req_sla_res_brch_reg,
        "object":fig_req_sla_res_brch_reg,
        "req_id":"req_sla_res_brch_reg",
        "purpose":p_req_sla_res_brch_reg,
        "benefit":b_req_sla_res_brch_reg,
        "chart_df": breach_counts_req_sla_res_brch_reg_df,
        "chart_type": "bar",
        "x_col": "Region",
        "y_col": "Breach Count",
        "value_col": None,
        "name_col": None
    })
    st.markdown("---")

    st.subheader("3.8 Resolution SLA Breach Rate by Priority")
    p_req_sla_res_brch_prio, b_req_sla_res_brch_prio = "Identify if priorities are prone to breaching Resolution SLAs.", "Pinpoint priorities needing process improvements."
    st.caption(f"**Purpose:** {p_req_sla_res_brch_prio} / **Benefit:** {b_req_sla_res_brch_prio}")

    # Local filters for req_sla_res_brch_prio (formerly new3)
    df_req_sla_res_brch_prio_local, lfilt_req_sla_res_brch_prio_desc, _, _, _ = apply_local_filters(
        df_filtered_global, 'req_sla_res_brch_prio', 'Opened',
        allow_time_slice=True, allow_period=True, allow_region=True, allow_account=True,
        allow_priority=False, allow_state=True, allow_category=True, allow_subcategory=True # Priority is the X-axis
    )
    local_filters_desc_map['req_sla_res_brch_prio'] = lfilt_req_sla_res_brch_prio_desc

    fig_req_sla_res_brch_prio = None; prio_summary_req_sla_res_brch_prio_df = None
    title_req_sla_res_brch_prio = "Resolution SLA Breach Rate by Priority"
    if 'Priority' in df_req_sla_res_brch_prio_local.columns and 'Target' in df_req_sla_res_brch_prio_local.columns and 'Has breached' in df_req_sla_res_brch_prio_local.columns and 'Task' in df_req_sla_res_brch_prio_local.columns:
        sla_prio_df_req_sla_res_brch_prio = df_req_sla_res_brch_prio_local[df_req_sla_res_brch_prio_local['Target']=='Resolution'].copy()
        if not sla_prio_df_req_sla_res_brch_prio.empty:
            sla_prio_df_req_sla_res_brch_prio['Is_Breached'] = sla_prio_df_req_sla_res_brch_prio['Has breached'].map({True:1,False:0,pd.NA:0, 'TRUE':1, 'FALSE':0}).astype(int)
            prio_summary_req_sla_res_brch_prio_df = sla_prio_df_req_sla_res_brch_prio.groupby('Priority').agg(Total_SLAs=('Task','nunique'),Breached_SLAs=('Is_Breached','sum')).reset_index()
            prio_summary_req_sla_res_brch_prio_df['Breach Rate (%)'] = prio_summary_req_sla_res_brch_prio_df.apply(lambda r: (r['Breached_SLAs']/r['Total_SLAs']*100) if r['Total_SLAs']>0 else 0, axis=1)
            try: prio_summary_req_sla_res_brch_prio_df['SortKey']=prio_summary_req_sla_res_brch_prio_df['Priority'].astype(str).str.extract(r'(\d+)').astype(int); prio_summary_req_sla_res_brch_prio_df=prio_summary_req_sla_res_brch_prio_df.sort_values('SortKey')
            except: prio_summary_req_sla_res_brch_prio_df=prio_summary_req_sla_res_brch_prio_df.sort_values('Priority')
            if not prio_summary_req_sla_res_brch_prio_df.empty:
                fig_req_sla_res_brch_prio=px.bar(prio_summary_req_sla_res_brch_prio_df,x='Priority',y='Breach Rate (%)',title=title_req_sla_res_brch_prio)
                fig_req_sla_res_brch_prio.update_layout(yaxis_title="Breach Rate (%)",yaxis_range=[0,105]).update_traces(texttemplate='%{y:.1f}%', textposition='outside', marker_color=COLOR_BREACH)
                display_chart_with_genai(fig_req_sla_res_brch_prio, prio_summary_req_sla_res_brch_prio_df, title_req_sla_res_brch_prio, p_req_sla_res_brch_prio, b_req_sla_res_brch_prio, "req_sla_res_brch_prio", chart_type="bar", x_col="Priority", y_col="Breach Rate (%)")
            else: st.info("No data for SLA breach rate by priority (filtered).")
        else: st.info("No 'Resolution' target tickets for SLA breach rate by priority.")
    else: st.info("Missing 'Priority', 'Target', 'Has breached', or 'Task'.")
    figures_for_export.append({
        "title":title_req_sla_res_brch_prio,
        "object":fig_req_sla_res_brch_prio,
        "req_id":"req_sla_res_brch_prio",
        "purpose":p_req_sla_res_brch_prio,
        "benefit":b_req_sla_res_brch_prio,
        "chart_df": prio_summary_req_sla_res_brch_prio_df,
        "chart_type": "bar",
        "x_col": "Priority",
        "y_col": "Breach Rate (%)",
        "value_col": None,
        "name_col": None
    })
    st.markdown("---")

    # --- 📈 Efficiency & Duration Analysis ---
    st.header("📈 Efficiency & Duration Analysis")

    st.subheader("4.1 Average Resolution Time by Region")
    p_add1, b_add1 = "Compare avg. ticket resolution time across regions.", "Highlights regional performance variations."
    st.caption(f"**Purpose:** {p_add1} / **Benefit:** {b_add1}")

    # Local filters for add1_eff
    df_add1_local, lfilt_add1_desc, _, _, _ = apply_local_filters(
        df_filtered_global, 'add1_eff', 'Opened',
        allow_time_slice=True, allow_period=True, allow_region=False, allow_account=True, # Region is the X-axis
        allow_priority=True, allow_state=True, allow_category=True, allow_subcategory=True
    )
    local_filters_desc_map['add1_eff'] = lfilt_add1_desc

    add_fig1 = None; avg_res_time_add1_df = None
    add_fig1_title = 'Average Resolution Time by Region'
    if 'Opened' in df_add1_local.columns and 'Resolved' in df_add1_local.columns and 'Region' in df_add1_local.columns:
        dur_df_res_add1 = df_add1_local.dropna(subset=['Opened','Resolved','Region']).copy()
        if not dur_df_res_add1.empty:
            dur_df_res_add1['Opened']=pd.to_datetime(dur_df_res_add1['Opened'],errors='coerce',dayfirst=('Opened' in DAY_FIRST_TRUE_COLS))
            dur_df_res_add1['Resolved']=pd.to_datetime(dur_df_res_add1['Resolved'],errors='coerce',dayfirst=('Resolved' in DAY_FIRST_TRUE_COLS))
            dur_df_res_add1=dur_df_res_add1.dropna(subset=['Opened','Resolved'])
            dur_df_res_add1['ResTime_Hours']=(dur_df_res_add1['Resolved']-dur_df_res_add1['Opened'])/pd.Timedelta(hours=1)
            dur_df_res_add1=dur_df_res_add1[dur_df_res_add1['ResTime_Hours']>=0]
            avg_res_time_add1_df = dur_df_res_add1.groupby('Region')['ResTime_Hours'].mean().reset_index()
            if not avg_res_time_add1_df.empty:
                add_fig1=px.bar(avg_res_time_add1_df,x='Region',y='ResTime_Hours',title=add_fig1_title,labels={'ResTime_Hours':'Avg. Res. Time (Hours)'},text_auto='.2f')
                add_fig1.update_layout(xaxis={'categoryorder':'total descending'}).update_traces(marker_color=px.colors.qualitative.Pastel[0],textposition='outside')
                display_chart_with_genai(add_fig1, avg_res_time_add1_df, add_fig1_title, p_add1, b_add1, "add1_eff", chart_type="bar", x_col="Region", y_col="ResTime_Hours")
            else: st.info("No valid resolution time data by region (filtered).")
        else: st.info("No data with Opened, Resolved, Region (filtered).")
    else: st.info("Missing 'Opened', 'Resolved', or 'Region'.")
    figures_for_export.append({
        "title":add_fig1_title,
        "object":add_fig1,
        "req_id":"add1_eff",
        "purpose":p_add1,
        "benefit":b_add1,
        "chart_df": avg_res_time_add1_df,
        "chart_type": "bar",
        "x_col": "Region",
        "y_col": "ResTime_Hours",
        "value_col": None,
        "name_col": None
    })
    st.markdown("---")

    st.subheader("4.2 Average Ticket Lifecycle Time by Region (Aging)")
    p_add2, b_add2 = "Compare avg. ticket lifecycle time (Open to Close) by region.", "Broader view of ticket handling efficiency."
    st.caption(f"**Purpose:** {p_add2} / **Benefit:** {b_add2}")

    # Local filters for add2_eff
    df_add2_local, lfilt_add2_desc, _, _, _ = apply_local_filters(
        df_filtered_global, 'add2_eff', 'Opened',
        allow_time_slice=True, allow_period=True, allow_region=False, allow_account=True, # Region is the X-axis
        allow_priority=True, allow_state=True, allow_category=True, allow_subcategory=True
    )
    local_filters_desc_map['add2_eff'] = lfilt_add2_desc

    add_fig2 = None; avg_life_time_add2_df = None
    add_fig2_title = 'Average Ticket Lifecycle Time by Region'
    if 'Opened' in df_add2_local.columns and 'Closed' in df_add2_local.columns and 'Region' in df_add2_local.columns:
        dur_df_life_add2 = df_add2_local.dropna(subset=['Opened','Closed','Region']).copy()
        if not dur_df_life_add2.empty:
            dur_df_life_add2['Opened']=pd.to_datetime(dur_df_life_add2['Opened'],errors='coerce',dayfirst=('Opened' in DAY_FIRST_TRUE_COLS))
            dur_df_life_add2['Closed']=pd.to_datetime(dur_df_life_add2['Closed'],errors='coerce',dayfirst=('Closed' in DAY_FIRST_TRUE_COLS))
            dur_df_life_add2=dur_df_life_add2.dropna(subset=['Opened','Closed'])
            dur_df_life_add2['LifeTime_Days']=(dur_df_life_add2['Closed']-dur_df_life_add2['Opened'])/pd.Timedelta(days=1)
            dur_df_life_add2=dur_df_life_add2[dur_df_life_add2['LifeTime_Days']>=0]
            avg_life_time_add2_df = dur_df_life_add2.groupby('Region')['LifeTime_Days'].mean().reset_index()
            if not avg_life_time_add2_df.empty:
                add_fig2=px.bar(avg_life_time_add2_df,x='Region',y='LifeTime_Days',title=add_fig2_title,labels={'LifeTime_Days':'Avg. Lifecycle Time (Days)'},text_auto='.1f')
                add_fig2.update_layout(xaxis={'categoryorder':'total descending'}).update_traces(marker_color=px.colors.qualitative.Pastel[1],textposition='outside')
                display_chart_with_genai(add_fig2, avg_life_time_add2_df, add_fig2_title, p_add2, b_add2, "add2_eff", chart_type="bar", x_col="Region", y_col="LifeTime_Days")
            else: st.info("No valid lifecycle time data by region (filtered).")
        else: st.info("No data with Opened, Closed, Region (filtered).")
    else: st.info("Missing 'Opened', 'Closed', or 'Region'.")
    figures_for_export.append({
        "title":add_fig2_title,
        "object":add_fig2,
        "req_id":"add2_eff",
        "purpose":p_add2,
        "benefit":b_add2,
        "chart_df": avg_life_time_add2_df,
        "chart_type": "bar",
        "x_col": "Region",
        "y_col": "LifeTime_Days",
        "value_col": None,
        "name_col": None
    })
    st.markdown("---")

    st.subheader("4.3 Resolution Time Distribution by SLA Status (Resolution Target)")
    p_add4, b_add4 = "Compare average resolution times for tickets that met vs. breached 'Resolution' SLA.", "Visualizes average impact of SLA breaches on duration."
    st.caption(f"**Purpose:** {p_add4} / **Benefit:** {b_add4}")

    # Local filters for add4_eff
    df_add4_local, lfilt_add4_desc, _, _, _ = apply_local_filters(
        df_filtered_global, 'add4_eff', 'Opened',
        allow_time_slice=True, allow_period=True, allow_region=True, allow_account=True,
        allow_priority=True, allow_state=True, allow_category=True, allow_subcategory=True
    )
    local_filters_desc_map['add4_eff'] = lfilt_add4_desc

    add_fig4 = None; avg_res_time_sla_df = None # Changed variable name for clarity
    add_fig4_title = "Average Resolution Time by Resolution SLA Status" # Updated title
    if 'Opened' in df_add4_local.columns and 'Resolved' in df_add4_local.columns and 'Has breached' in df_add4_local.columns and 'Target' in df_add4_local.columns:
        res_sla_df_base_add4 = df_add4_local[df_add4_local['Target']=='Resolution'].dropna(subset=['Opened','Resolved','Has breached']).copy()
        if not res_sla_df_base_add4.empty:
            res_sla_df_base_add4['Opened']=pd.to_datetime(res_sla_df_base_add4['Opened'],errors='coerce',dayfirst=('Opened' in DAY_FIRST_TRUE_COLS))
            res_sla_df_base_add4['Resolved']=pd.to_datetime(res_sla_df_base_add4['Resolved'],errors='coerce',dayfirst=('Resolved' in DAY_FIRST_TRUE_COLS))
            res_sla_df_base_add4=res_sla_df_base_add4.dropna(subset=['Opened','Resolved'])
            
            # Calculate Resolution Time in Hours
            res_sla_df_base_add4['ResTime_Hours']=(res_sla_df_base_add4['Resolved']-res_sla_df_base_add4['Opened'])/pd.Timedelta(hours=1)
            res_sla_df_base_add4=res_sla_df_base_add4[res_sla_df_base_add4['ResTime_Hours']>=0] # Filter out negative durations

            # Map 'Has breached' to 'SLA Status' for display
            res_sla_df_base_add4['SLA Status']=res_sla_df_base_add4['Has breached'].map({True:'Breached',False:'Met',pd.NA:'Unknown', 'TRUE':'Breached', 'FALSE':'Met'})
            res_sla_df_base_add4=res_sla_df_base_add4[res_sla_df_base_add4['SLA Status']!='Unknown'] # Exclude unknown statuses

            if not res_sla_df_base_add4.empty:
                # Group by SLA Status and calculate mean resolution time for the bar chart
                avg_res_time_sla_df = res_sla_df_base_add4.groupby('SLA Status')['ResTime_Hours'].mean().reset_index()
                
                # Define a specific order for 'Met' and 'Breached' and assign colors
                status_order = ['Met', 'Breached']
                avg_res_time_sla_df['SLA Status'] = pd.Categorical(avg_res_time_sla_df['SLA Status'], categories=status_order, ordered=True)
                avg_res_time_sla_df = avg_res_time_sla_df.sort_values('SLA Status')

                # Define custom colors: Green for Met, Red for Breached
                color_map = {'Met': COLOR_SLA_RESPONSE, 'Breached': COLOR_BREACH} # Using COLOR_BREACH for consistency

                if not avg_res_time_sla_df.empty:
                    add_fig4=px.bar(avg_res_time_sla_df,
                                    x='SLA Status',
                                    y='ResTime_Hours',
                                    color='SLA Status', # Use color to differentiate status
                                    color_discrete_map=color_map, # Apply custom colors
                                    title=add_fig4_title,
                                    labels={'SLA Status':'Resolution SLA Status','ResTime_Hours':'Avg. Resolution Time (Hours)'},
                                    text_auto='.2f' # Display average with 2 decimal places on bars
                                   )
                    add_fig4.update_layout(yaxis_title="Avg. Resolution Time (Hours)", showlegend=False) # Hide legend as color indicates status
                    display_chart_with_genai(add_fig4, avg_res_time_sla_df, add_fig4_title, p_add4, b_add4, "add4_eff", chart_type="bar", x_col="SLA Status", y_col="ResTime_Hours")
                else: 
                    st.info("No aggregated data for Average Resolution Time vs SLA Status (filtered).")
            else: 
                st.info("No valid data for Resolution Time vs SLA Status after filtering (e.g., no resolution times or SLA statuses).")
        else: 
            st.info("No data with 'Opened', 'Resolved', 'Has breached' for Target='Resolution' (filtered).")
    else: 
        st.info("Missing one or more required columns ('Opened', 'Resolved', 'Has breached', 'Target') for Resolution Time vs SLA Status analysis.")
    figures_for_export.append({
        "title":add_fig4_title,
        "object":add_fig4,
        "req_id":"add4_eff",
        "purpose":p_add4,
        "benefit":b_add4,
        "chart_df": avg_res_time_sla_df, # Pass the aggregated df for better LLM summary
        "chart_type": "bar", # Changed to bar
        "x_col": "SLA Status",
        "y_col": "ResTime_Hours",
        "value_col": None,
        "name_col": None
    })
    st.markdown("---")

    st.subheader("4.4 Average Resolution Time by Priority (MTTR)")
    p_new2, b_new2 = "Analyze if higher priority tickets are resolved faster.", "Validates priority handling effectiveness."
    st.caption(f"**Purpose:** {p_new2} / **Benefit:** {b_new2}")

    # Local filters for new2_eff
    df_new2_local, lfilt_new2_desc, _, _, _ = apply_local_filters(
        df_filtered_global, 'new2_eff', 'Opened',
        allow_time_slice=True, allow_period=True, allow_region=True, allow_account=True,
        allow_priority=False, allow_state=True, allow_category=True, allow_subcategory=True # Priority is the X-axis
    )
    local_filters_desc_map['new2_eff'] = lfilt_new2_desc

    fig_new2 = None; avg_res_prio_new2_df = None
    fig_new2_title = "Average Resolution Time by Priority"
    if 'Priority' in df_new2_local.columns and 'Opened' in df_new2_local.columns and 'Resolved' in df_new2_local.columns:
        prio_res_df_new2 = df_new2_local.dropna(subset=['Priority','Opened','Resolved']).copy()
        if not prio_res_df_new2.empty:
            prio_res_df_new2['Opened']=pd.to_datetime(prio_res_df_new2['Opened'],errors='coerce',dayfirst=('Opened' in DAY_FIRST_TRUE_COLS))
            prio_res_df_new2['Resolved']=pd.to_datetime(prio_res_df_new2['Resolved'],errors='coerce',dayfirst=('Resolved' in DAY_FIRST_TRUE_COLS))
            prio_res_df_new2=prio_res_df_new2.dropna(subset=['Opened','Resolved'])
            prio_res_df_new2['ResTime_Hours']=(prio_res_df_new2['Resolved']-prio_res_df_new2['Opened'])/pd.Timedelta(hours=1)
            prio_res_df_new2=prio_res_df_new2[prio_res_df_new2['ResTime_Hours']>=0]
            avg_res_prio_new2_df = prio_res_df_new2.groupby('Priority')['ResTime_Hours'].mean().reset_index()
            try: avg_res_prio_new2_df['SortKey']=avg_res_prio_new2_df['Priority'].astype(str).str.extract(r'(\d+)').astype(int); avg_res_prio_new2_df=avg_res_prio_new2_df.sort_values('SortKey')
            except: avg_res_prio_new2_df=avg_res_prio_new2_df.sort_values('Priority')
            if not avg_res_prio_new2_df.empty:
                fig_new2=px.bar(avg_res_prio_new2_df,x='Priority',y='ResTime_Hours',title=fig_new2_title,text_auto='.2f')
                fig_new2.update_layout(yaxis_title="Avg. Resolution Time (Hours)")
                display_chart_with_genai(fig_new2, avg_res_prio_new2_df, fig_new2_title, p_new2, b_new2, "new2_eff", chart_type="bar", x_col="Priority", y_col="ResTime_Hours")
            else: st.info("No valid data for avg. res. time by priority (filtered).")
        else: st.info("No data with Priority, Opened, Resolved (filtered).")
    else: st.info("Missing 'Priority', 'Opened', or 'Resolved'.")
    figures_for_export.append({
        "title":fig_new2_title,
        "object":fig_new2,
        "req_id":"new2_eff",
        "purpose":p_new2,
        "benefit":b_new2,
        "chart_df": avg_res_prio_new2_df,
        "chart_type": "bar",
        "x_col": "Priority",
        "y_col": "ResTime_Hours",
        "value_col": None,
        "name_col": None
    })
    st.markdown("---")

    # --- 🏷️ Issue Categorization (Conditional) ---
    if 'Category' in df_filtered_global.columns:
        st.header("🏷️ Issue Categorization")
        st.subheader("5.1 Ticket Distribution by Category")
        p_new6a,b_new6a = "Understand common issue types by primary category.","Drives problem mgt, aids solution development."
        st.caption(f"**Purpose:** {p_new6a} / **Benefit:** {b_new6a}")

        # Local filters for new6a
        df_new6a_local, lfilt_new6a_desc, _, _, _ = apply_local_filters(
            df_filtered_global, 'new6a', 'Opened',
            allow_time_slice=True, allow_period=True, allow_region=True, allow_account=True,
            allow_priority=True, allow_state=True, allow_category=False, allow_subcategory=True # Category is the X-axis
        )
        local_filters_desc_map['new6a'] = lfilt_new6a_desc

        fig_new6a = None; cat_counts_new6a_df = None
        fig_new6a_title = "Ticket Distribution by Category (Top 15)"
        if 'Task' in df_new6a_local.columns and 'Category' in df_new6a_local.columns:
            cat_counts_new6a_df = df_new6a_local.groupby('Category')['Task'].nunique().nlargest(15).reset_index(name='Ticket Count')
            if not cat_counts_new6a_df.empty:
                fig_new6a = px.bar(cat_counts_new6a_df,x='Category',y='Ticket Count',title=fig_new6a_title,text_auto=True)
                fig_new6a.update_layout(xaxis_title="Category",yaxis_title="Number of Tickets",xaxis={'categoryorder':'total descending'})
                display_chart_with_genai(fig_new6a, cat_counts_new6a_df, fig_new6a_title, p_new6a, b_new6a, "new6a", chart_type="bar", x_col="Category", y_col="Ticket Count")
            else: st.info("No ticket data by category (filtered).")
        else: st.info("Missing 'Task' or 'Category' for category distribution.")
        figures_for_export.append({
            "title":fig_new6a_title,
            "object":fig_new6a,
            "req_id":"new6a",
            "purpose":p_new6a,
            "benefit":b_new6a,
            "chart_df": cat_counts_new6a_df,
            "chart_type": "bar",
            "x_col": "Category",
            "y_col": "Ticket Count",
            "value_col": None,
            "name_col": None
        })
        st.markdown("---")

        # NEW: Response tickets count by Category within section 5.1
        st.markdown("##### 5.1.1 'Response' Target Tickets by Category")
        p_new6a_resp_count = "To show the distribution of tickets with 'Response' SLA targets across categories."
        b_new6a_resp_count = "Helps identify which categories frequently trigger Response SLAs, aiding targeted process improvements."
        st.caption(f"**Purpose:** {p_new6a_resp_count} / **Benefit:** {b_new6a_resp_count}")

        fig_new6a_resp = None; resp_cat_counts_df = None
        title_new6a_resp = "Response Target Tickets by Category (Top 15)"
        if 'Task' in df_new6a_local.columns and 'Category' in df_new6a_local.columns and 'Target' in df_new6a_local.columns:
            resp_df_temp = df_new6a_local[df_new6a_local['Target'] == 'Response'].copy()
            if not resp_df_temp.empty:
                resp_cat_counts_df = resp_df_temp.groupby('Category')['Task'].nunique().nlargest(15).reset_index(name='Response Ticket Count')
                if not resp_cat_counts_df.empty:
                    fig_new6a_resp = px.bar(resp_cat_counts_df, x='Category', y='Response Ticket Count', 
                                            title=title_new6a_resp, text_auto=True,
                                            color_discrete_sequence=[COLOR_SLA_RESPONSE])
                    fig_new6a_resp.update_layout(xaxis_title="Category", yaxis_title="Number of Response Tickets", 
                                                 xaxis={'categoryorder':'total descending'})
                    display_chart_with_genai(fig_new6a_resp, resp_cat_counts_df, title_new6a_resp, p_new6a_resp_count, b_new6a_resp_count, "new6a_resp_count", chart_type="bar", x_col="Category", y_col="Response Ticket Count")
                else: st.info("No 'Response' target tickets found by category (filtered).")
            else: st.info("No 'Response' target tickets in the filtered data.")
        else: st.info("Missing 'Task', 'Category', or 'Target' for Response ticket category distribution.")
        figures_for_export.append({
            "title": title_new6a_resp,
            "object": fig_new6a_resp,
            "req_id": "new6a_resp_count",
            "purpose": p_new6a_resp_count,
            "benefit": b_new6a_resp_count,
            "chart_df": resp_cat_counts_df,
            "chart_type": "bar",
            "x_col": "Category",
            "y_col": "Response Ticket Count",
            "value_col": None,
            "name_col": None
        })
        st.markdown("---")

        if 'Subcategory' in df_filtered_global.columns:
            st.subheader("5.2 Ticket Distribution by Subcategory")
            p_new6b,b_new6b = "Drill down into subcategories for a selected category.","Precise identification of problem areas."
            st.caption(f"**Purpose:** {p_new6b} / **Benefit:** {b_new6b}")

            # Local filters for new6b (Note: Category filter within apply_local_filters is used for the dropdown list)
            df_new6b_local, lfilt_new6b_desc, _, _, _ = apply_local_filters(
                df_filtered_global, 'new6b', 'Opened',
                allow_time_slice=True, allow_period=True, allow_region=True, allow_account=True,
                allow_priority=True, allow_state=True, allow_category=True, allow_subcategory=False # Subcategory is the X-axis
            )
            local_filters_desc_map['new6b_all_cat'] = lfilt_new6b_desc # Default ID for subcat chart, will be overwritten if specific cat selected
            
            fig_new6b = None; subcat_counts_new6b_df = None
            fig_new6b_title = "Subcategory Distribution" # Default title
            req_id_new6b = "new6b_all_cat" # Default for 'All' or if no category selection happens
            if 'Task' in df_new6b_local.columns and 'Subcategory' in df_new6b_local.columns:
                
                # Retrieve the locally filtered category from the `apply_local_filters` function
                # This needs to be managed separately as the category filter is part of the subcategory chart's local filters
                current_selected_category_for_subcat = st.session_state.get(f'new6b_cat', 'All Categories')
                
                df_subcat_for_chart = df_new6b_local.copy()
                if current_selected_category_for_subcat != 'All Categories' and 'Category' in df_subcat_for_chart.columns: 
                    df_subcat_for_chart = df_subcat_for_chart[df_subcat_for_chart['Category']==current_selected_category_for_subcat]
                    req_id_new6b = f"new6b_{current_selected_category_for_subcat.replace(' ','_').replace('/','_')}" 
                
                subcat_counts_new6b_df = df_subcat_for_chart.groupby('Subcategory')['Task'].nunique().nlargest(10).reset_index(name='Ticket Count')
                fig_new6b_title = f"Subcategory Distribution (Top 10 for '{current_selected_category_for_subcat}')"
                if not subcat_counts_new6b_df.empty:
                    fig_new6b = px.bar(subcat_counts_new6b_df,x='Subcategory',y='Ticket Count',title=fig_new6b_title,text_auto=True)
                    fig_new6b.update_layout(xaxis_title="Subcategory",yaxis_title="Tickets",xaxis={'categoryorder':'total descending'})
                    display_chart_with_genai(fig_new6b, subcat_counts_new6b_df, fig_new6b_title, p_new6b, b_new6b, req_id_new6b, chart_type="bar", x_col="Subcategory", y_col="Ticket Count")
                else: st.info(f"No subcategory data for '{current_selected_category_for_subcat}' (filtered).")
            else: st.info("Missing 'Task' or 'Subcategory' for subcategory distribution.")
            figures_for_export.append({
                "title":fig_new6b_title,
                "object":fig_new6b,
                "req_id":req_id_new6b, # Use the dynamically created req_id
                "purpose":p_new6b,
                "benefit":b_new6b,
                "chart_df": subcat_counts_new6b_df,
                "chart_type": "bar",
                "x_col": "Subcategory",
                "y_col": "Ticket Count",
                "value_col": None,
                "name_col": None
            })
            st.markdown("---")
    else: 
        st.sidebar.info("Optional column 'Category' not found. 'Issue Categorization' section disabled.")
        st.info("Optional column 'Category' not found. 'Issue Categorization' section not available.")


    # --- 🔍 Advanced Analytics & Trends ---
    st.header("🔍 Advanced Analytics & Trends")

    st.subheader("6.1 Open Ticket Backlog Trend")
    p_new4, b_new4 = "Monitor unresolved (open) tickets over time.", "Key indicator of service desk health."
    st.caption(f"**Purpose:** {p_new4} / **Benefit:** {b_new4}")

    # Local filters for new4_adv
    # For backlog, 'Opened' is the primary date column for time slicing
    df_new4_local, lfilt_new4_desc, _, selected_time_slice_new4, freq_new4 = apply_local_filters(
        df_filtered_global, 'new4_adv', 'Opened', 
        allow_time_slice=True, allow_period=False, # Period is handled by the loop/snapshot logic, not a single selectbox
        allow_region=True, allow_account=True, allow_priority=True, allow_state=True, allow_category=True, allow_subcategory=True
    )
    local_filters_desc_map['new4_adv'] = lfilt_new4_desc

    fig_new4 = None; trend_df_new4 = None
    fig_new4_title = f"Open Ticket Backlog Trend ({selected_time_slice_new4}ly Snapshots)" # Dynamic title based on local filter
    if 'Opened' in df_new4_local.columns and ('Closed' in df_new4_local.columns or 'State' in df_new4_local.columns) and 'Task' in df_new4_local.columns:
        backlog_df_new4 = df_new4_local.copy()
        
        # Ensure OpenedDate is UTC-aware and drop NaT
        backlog_df_new4['OpenedDate'] = pd.to_datetime(backlog_df_new4['Opened'], errors='coerce', dayfirst=('Opened' in DAY_FIRST_TRUE_COLS))
        if backlog_df_new4['OpenedDate'].dt.tz is None:
            backlog_df_new4['OpenedDate'] = backlog_df_new4['OpenedDate'].dt.tz_localize('UTC', ambiguous='NaT', nonexistent='NaT')
        else:
            backlog_df_new4['OpenedDate'] = backlog_df_new4['OpenedDate'].dt.tz_convert('UTC')
        backlog_df_new4 = backlog_df_new4.dropna(subset=['OpenedDate'])
        
        # Ensure ClosedDate is UTC-aware if present and drop NaT (only for this column)
        if 'Closed' in backlog_df_new4.columns and pd.api.types.is_datetime64_any_dtype(backlog_df_new4['Closed']):
            backlog_df_new4['ClosedDate'] = pd.to_datetime(backlog_df_new4['Closed'], errors='coerce', dayfirst=('Closed' in DAY_FIRST_TRUE_COLS))
            if backlog_df_new4['ClosedDate'].dt.tz is None:
                backlog_df_new4['ClosedDate'] = backlog_df_new4['ClosedDate'].dt.tz_localize('UTC', ambiguous='NaT', nonexistent='NaT')
            else:
                backlog_df_new4['ClosedDate'] = backlog_df_new4['ClosedDate'].dt.tz_convert('UTC')
        else: # Ensure 'ClosedDate' column exists, even if all NaT, for consistent logic below
            backlog_df_new4['ClosedDate'] = pd.NaT # Fill with NaT if column missing or not datetime-like

        if not backlog_df_new4.empty:
            min_dt_for_range = backlog_df_new4['OpenedDate'].min()
            
            # Aggregate all possible end dates to determine the max range
            max_dt_values = [backlog_df_new4['OpenedDate'].max()]
            if 'ClosedDate' in backlog_df_new4.columns and not backlog_df_new4['ClosedDate'].isnull().all():
                max_dt_values.append(backlog_df_new4['ClosedDate'].max(skipna=True))
            max_dt_values.append(datetime.now(timezone.utc)) # Always include current time for backlog to reflect up-to-date status

            # Filter out NaT values before finding the max to avoid errors
            max_dt_for_range = pd.Series([d for d in max_dt_values if pd.notna(d)]).max() 
            
            # Check for valid date range
            if pd.isna(min_dt_for_range) or pd.isna(max_dt_for_range) or min_dt_for_range > max_dt_for_range:
                st.info("Not enough valid date range for backlog trend (min/max dates invalid).")
                fig_new4 = None
                trend_df_new4 = None
                # No `continue` needed here as we are not in a loop, but need to exit this `if` block.
            else:
                # Calculate start and end of periods for the date range
                start_of_first_period_naive = min_dt_for_range.to_period(freq_new4).start_time
                end_of_last_period_naive = max_dt_for_range.to_period(freq_new4).end_time

                # Ensure these are UTC-aware before passing to date_range
                start_of_first_period_utc = start_of_first_period_naive.tz_localize('UTC')
                end_of_last_period_utc = end_of_last_period_naive.tz_localize('UTC')
                
                date_rng_new4 = pd.date_range(start_of_first_period_utc, end_of_last_period_utc, freq=freq_new4)
                
                # Fallback for empty date_range if start/end are in the same period but range somehow yields empty (unlikely but safe)
                if date_rng_new4.empty and start_of_first_period_utc <= end_of_last_period_utc:
                    date_rng_new4 = pd.DatetimeIndex([end_of_last_period_utc])
                elif date_rng_new4.empty: # Truly empty and invalid range
                    st.info("Could not determine valid date range for backlog trend (generated range is empty).")
                    fig_new4 = None
                    trend_df_new4 = None
                else: # Valid date range, proceed with backlog calculation
                    backlog_trend_data = []
                    for period_end_utc in date_rng_new4:
                        # Tickets opened up to or on this period end
                        opened_up_to_period = backlog_df_new4[backlog_df_new4['OpenedDate'] <= period_end_utc].copy() # Use .copy() to avoid SettingWithCopyWarning
                        
                        currently_open_count = 0
                        if not opened_up_to_period.empty:
                            is_open_mask = pd.Series(True, index=opened_up_to_period.index) # Assume open by default
                            
                            has_state_col = 'State' in opened_up_to_period.columns
                            has_closed_date_col = 'ClosedDate' in opened_up_to_period.columns and \
                                                  pd.api.types.is_datetime64_any_dtype(opened_up_to_period['ClosedDate'])

                            if has_state_col and has_closed_date_col:
                                # A ticket is considered closed at period_end_utc if:
                                # (its state is in CLOSED_STATES_LIST) AND (it has a valid ClosedDate AND that ClosedDate is <= period_end_utc)
                                is_closed_at_period_end = (opened_up_to_period['State'].isin(CLOSED_STATES_LIST)) & \
                                                          (opened_up_to_period['ClosedDate'].notna()) & \
                                                          (opened_up_to_period['ClosedDate'] <= period_end_utc)
                                is_open_mask = ~is_closed_at_period_end
                            elif has_state_col:
                                # If no ClosedDate, rely only on State. Tickets in closed states are considered closed.
                                is_open_mask = ~opened_up_to_period['State'].isin(CLOSED_STATES_LIST)
                            elif has_closed_date_col:
                                # If no State, rely only on ClosedDate. Tickets with valid ClosedDate <= period_end_utc are closed.
                                is_closed_at_period_end = (opened_up_to_period['ClosedDate'].notna()) & \
                                                          (opened_up_to_period['ClosedDate'] <= period_end_utc)
                                is_open_mask = ~is_closed_at_period_end
                            # If neither has_state_col nor has_closed_date_col are true, then `is_open_mask` remains True (all opened are open)
                            
                            currently_open_count = opened_up_to_period[is_open_mask]['Task'].nunique()
                        
                        backlog_trend_data.append({
                            'Period': format_period_label(period_end_utc, freq_new4), 
                            'Open Backlog': currently_open_count
                        })
                    
                    if backlog_trend_data:
                        trend_df_new4 = pd.DataFrame(backlog_trend_data)
                        # Convert 'Period' to datetime for proper sorting in Plotly line chart
                        # Use the period_end_utc directly for consistent sorting, then format for display
                        trend_df_new4['Sort_Period'] = date_rng_new4 # Matches the list order
                        trend_df_new4 = trend_df_new4.sort_values('Sort_Period')
                        
                        if not trend_df_new4.empty:
                            fig_new4 = px.line(trend_df_new4,x='Period',y='Open Backlog',title=fig_new4_title,markers=True)
                            fig_new4.update_layout(yaxis_title="Number of Open Tickets")
                            display_chart_with_genai(fig_new4, trend_df_new4, fig_new4_title, p_new4, b_new4, "new4_adv", chart_type="line", x_col="Period", y_col="Open Backlog")
                        else: st.info("Could not generate backlog trend (filtered data resulted in empty trend DataFrame).")
                    else: st.info("Could not generate backlog trend data (no periods found).")
        else: st.info("No tickets with valid 'Opened' dates for backlog after timezone handling.")
    else: st.info("Missing 'Opened', ('Closed' or 'State'), or 'Task' columns for backlog trend.")
    figures_for_export.append({
        "title":fig_new4_title,
        "object":fig_new4,
        "req_id":"new4_adv",
        "purpose":p_new4,
        "benefit":b_new4,
        "chart_df": trend_df_new4,
        "chart_type": "line",
        "x_col": "Period",
        "y_col": "Open Backlog",
        "value_col": None,
        "name_col": None
    })
    st.markdown("---")

    st.subheader("6.2 Heatmap of Ticket Creation (Day of Week vs. Hour of Day)")
    p_new5,b_new5 = "Identify peak times/days for ticket submission.","Aids resource planning, shows ticket patterns."
    st.caption(f"**Purpose:** {p_new5} / **Benefit:** {b_new5}")

    # Local filters for new5_adv
    df_new5_local, lfilt_new5_desc, _, _, _ = apply_local_filters(
        df_filtered_global, 'new5_adv', 'Opened',
        allow_time_slice=True, allow_period=True, allow_region=True, allow_account=True,
        allow_priority=True, allow_state=True, allow_category=True, allow_subcategory=True
    )
    local_filters_desc_map['new5_adv'] = lfilt_new5_desc

    fig_new5 = None; pivot_hm_new5_df = None
    fig_new5_title = "Heatmap of Ticket Creation Times"
    if 'Opened' in df_new5_local.columns and 'Task' in df_new5_local.columns:
        heatmap_df_new5 = df_new5_local.dropna(subset=['Opened']).copy()
        heatmap_df_new5['OpenedDT'] = pd.to_datetime(heatmap_df_new5['Opened'], errors='coerce', dayfirst=('Opened' in DAY_FIRST_TRUE_COLS))
        heatmap_df_new5 = heatmap_df_new5.dropna(subset=['OpenedDT'])
        if not heatmap_df_new5.empty:
            heatmap_df_new5['DayOfWeek'] = heatmap_df_new5['OpenedDT'].dt.day_name()
            heatmap_df_new5['HourOfDay'] = heatmap_df_new5['OpenedDT'].dt.hour
            ticket_hm_new5 = heatmap_df_new5.groupby(['DayOfWeek','HourOfDay'])['Task'].nunique().reset_index(name='Ticket Count')
            try:
                days_order = ["Monday","Tuesday","Wednesday","Thursday","Friday","Saturday","Sunday"]
                pivot_hm_new5_df = ticket_hm_new5.pivot_table(index='DayOfWeek',columns='HourOfDay',values='Ticket Count',fill_value=0)
                pivot_hm_new5_df = pivot_hm_new5_df.reindex(index=days_order).fillna(0) # Reindex to ensure consistent day order
                for hr in range(24): # Ensure all hours are present, even if no tickets
                    if hr not in pivot_hm_new5_df.columns: pivot_hm_new5_df[hr]=0
                pivot_hm_new5_df = pivot_hm_new5_df[sorted(pivot_hm_new5_df.columns)] 
                if not pivot_hm_new5_df.empty:
                    fig_new5 = px.imshow(pivot_hm_new5_df,labels=dict(x="Hour of Day",y="Day of Week",color="Ticket Count"),title=fig_new5_title,color_continuous_scale=px.colors.sequential.Viridis)
                    fig_new5.update_xaxes(side="bottom",tickmode='array',tickvals=list(range(24)),ticktext=[str(h) for h in range(24)])
                    display_chart_with_genai(fig_new5, pivot_hm_new5_df, fig_new5_title, p_new5, b_new5, "new5_adv", chart_type="heatmap", x_col="HourOfDay", y_col="DayOfWeek", value_col="Ticket Count") # Pass value_col for heatmap summary
                else: st.info("Could not generate heatmap pivot (filtered).")
            except Exception as e_hm: st.warning(f"Could not generate heatmap: {e_hm}")
        else: st.info("No tickets with valid 'Opened' for heatmap.")
    else: st.info("Missing 'Opened' or 'Task'.")
    figures_for_export.append({
        "title":fig_new5_title,
        "object":fig_new5,
        "req_id":"new5_adv",
        "purpose":p_new5,
        "benefit":b_new5,
        "chart_df": pivot_hm_new5_df,
        "chart_type": "heatmap",
        "x_col": "HourOfDay",
        "y_col": "DayOfWeek",
        "value_col": "Ticket Count", # Explicitly pass for summary generation
        "name_col": None
    })
    st.markdown("---")

    # --- Export Buttons ---
    st.sidebar.markdown("---")
    st.sidebar.subheader("Export Report")

    if st.sidebar.button("Generate & Download Word Report", key="export_docx_main_final"):
        if not figures_for_export and not kpi_data: # Check if there's anything to export
            st.sidebar.warning("No figures or KPI data were generated to include in the Word report.")
        else:
            # Generate all GenAI summaries sequentially
            all_summaries = generate_all_chart_genai_summaries(figures_for_export)

            valid_figures_to_export = [item for item in figures_for_export if item.get("object") is not None]
            
            if not valid_figures_to_export and not kpi_data and not all_summaries:
                 st.sidebar.warning("No data, charts, or GenAI summaries to export.")
            else:
                with st.spinner("Compiling Word document..."):
                    try:
                        docx_bytes = create_export_docx(
                            kpi_data,
                            valid_figures_to_export, # Send only valid figures
                            global_filters_description,
                            local_filters_desc_map,
                            all_summaries # Pass the pre-generated summaries
                        )
                        st.sidebar.download_button(
                            label="Download Word Document (.docx)",
                            data=docx_bytes,
                            file_name=f"ServiceNow_Advanced_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        st.sidebar.success("Word document ready for download!")
                    except ImportError as ie:
                        if 'docx' in str(ie).lower():
                            st.sidebar.error(f"Word generation library missing: {ie}. Please install it: pip install python-docx")
                        elif 'kaleido' in str(ie).lower():
                             st.sidebar.error(f"Image export library 'kaleido' missing. Please install it: pip install kaleido")
                        else:
                            st.sidebar.error(f"An import error occurred: {ie}. Ensure all dependencies are installed.")
                        st.exception(ie)
                    except Exception as e:
                        st.sidebar.error(f"Error generating Word document: {e}")
                        st.exception(e)

elif df_original is not None and df_original.empty:
     st.warning("The uploaded file is empty or contains no data after initial processing.")
else:
    st.info("Upload an Excel file to begin analysis.")
    keys_to_clear_on_no_file = ['data', 'uploaded_file_name', 'global_periods', 'global_regions', 'global_accounts', 'current_freq']
    clear_genai_summaries_session_state() # Clear all GenAI summaries
    for key in list(st.session_state.keys()):
        if key.startswith("genai_summary_"): del st.session_state[key]
    for key in keys_to_clear_on_no_file: st.session_state.pop(key, None)

# --- Footer ---
st.markdown("---")
st.caption("Company Confidential. Report generated for internal analysis purposes. AI insights are experimental and may require review.")
